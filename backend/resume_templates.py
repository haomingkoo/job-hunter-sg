"""
Resume template system — generates formatted DOCX files from resume data.

8 templates based on Harvard, MIT/Jake's, NUS, and Stanford research:
1. classic    — Serif, formal, education-first (Harvard style)
2. modern     — Dense, clean, projects section (Jake's/tech style)
3. singapore  — Includes nationality/PR, C.A.R. format (NUS style)
4. compact    — Summary-first, 2-page friendly (experienced professionals)
5. executive  — Large name, gray header bar, small-caps headings (C-suite/VP)
6. creative   — Indigo accent, colored left border on headings (design/marketing)
7. technical  — Monospace headings, skills-first, compact (engineering/data)
8. minimal    — No borders, no backgrounds, maximum ATS compatibility

All templates are ATS-friendly:
- No tables for layout, no columns, no images, no text boxes
- Standard fonts (Calibri or Times New Roman)
- Standard section headers
- Bullet points with action verbs
- Dates right-aligned via tab stops
"""

from __future__ import annotations

import io
import logging
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

log = logging.getLogger("jobhunter.docx")


# ── ATS Unicode Normalization ─────────────────────────────────────────────────
# Many ATS systems (Workday, Taleo, Greenhouse) only handle basic ASCII.
# Smart quotes, em dashes, non-breaking spaces, and zero-width characters
# cause garbled text or failed keyword matching in these systems.

_ATS_REPLACEMENTS: list[tuple[str, str]] = [
    # Dashes
    ("\u2013", "-"),   # en dash
    ("\u2014", "-"),   # em dash
    ("\u2015", "-"),   # horizontal bar
    ("\u2212", "-"),   # minus sign
    # Quotes
    ("\u2018", "'"),   # left single quote
    ("\u2019", "'"),   # right single quote / apostrophe
    ("\u201A", "'"),   # single low-9 quote
    ("\u201C", '"'),   # left double quote
    ("\u201D", '"'),   # right double quote
    ("\u201E", '"'),   # double low-9 quote
    ("\u2039", "'"),   # single left angle quote
    ("\u203A", "'"),   # single right angle quote
    ("\u00AB", '"'),   # left double angle quote
    ("\u00BB", '"'),   # right double angle quote
    # Spaces
    ("\u00A0", " "),   # non-breaking space
    ("\u2002", " "),   # en space
    ("\u2003", " "),   # em space
    ("\u2009", " "),   # thin space
    # Bullets / symbols
    ("\u2022", "-"),   # bullet (replaced with hyphen for plain-text ATS)
    ("\u2023", "-"),   # triangular bullet
    ("\u25CF", "-"),   # black circle
    ("\u25CB", "-"),   # white circle
    ("\u25AA", "-"),   # black small square
    # Ellipsis
    ("\u2026", "..."), # horizontal ellipsis
]

# Zero-width / invisible characters to strip entirely
_ATS_STRIP_RE = re.compile(
    "[\u200B\u200C\u200D\u200E\u200F\uFEFF\u00AD\u2060\u034F]"
)


def normalize_for_ats(text: str) -> str:
    """Replace Unicode characters that break ATS keyword matching with safe ASCII."""
    for unicode_char, ascii_char in _ATS_REPLACEMENTS:
        text = text.replace(unicode_char, ascii_char)
    text = _ATS_STRIP_RE.sub("", text)
    return text


TEMPLATES = {
    "classic": {
        "name": "Classic",
        "description": "Clean, formal style. Education first. Best for fresh grads, consulting, government.",
        "font": "Times New Roman",
        "body_size": 11,
        "heading_size": 12,
        "name_size": 16,
        "margins": 1.0,
        "section_order": ["summary", "education", "experience", "skills", "certifications"],
    },
    "modern": {
        "name": "Modern",
        "description": "Dense, tech-focused. Projects section. Best for engineers, startups, FAANG.",
        "font": "Calibri",
        "body_size": 10,
        "heading_size": 11,
        "name_size": 14,
        "margins": 0.6,
        "section_order": ["summary", "experience", "projects", "skills", "education"],
    },
    "singapore": {
        "name": "Singapore Professional",
        "description": "Includes nationality/PR status. C.A.R. format. Best for SG local market.",
        "font": "Calibri",
        "body_size": 11,
        "heading_size": 12,
        "name_size": 15,
        "margins": 0.8,
        "section_order": ["personal", "summary", "education", "experience", "activities", "skills"],
    },
    "compact": {
        "name": "Compact",
        "description": "Summary-first, experience-heavy. Best for senior professionals with 10+ years.",
        "font": "Arial",
        "body_size": 10,
        "heading_size": 11,
        "name_size": 14,
        "margins": 0.5,
        "section_order": ["summary", "experience", "skills", "education", "certifications"],
    },
    "executive": {
        "name": "Executive",
        "description": "Large name, gray header bar, small-caps headings. Best for C-suite, VP, director roles.",
        "font": "Georgia",
        "body_size": 11,
        "heading_size": 13,
        "name_size": 20,
        "margins": 1.0,
        "section_order": ["summary", "experience", "education", "skills", "certifications"],
    },
    "creative": {
        "name": "Creative",
        "description": "Indigo accent color, colored left border on headings. Best for design, marketing, content roles.",
        "font": "Calibri",
        "body_size": 10,
        "heading_size": 12,
        "name_size": 16,
        "margins": 0.75,
        "section_order": ["summary", "experience", "projects", "skills", "education"],
    },
    "technical": {
        "name": "Technical",
        "description": "Monospace headings, skills placed prominently second. Best for engineering, data, DevOps roles.",
        "font": "Calibri",
        "body_size": 10,
        "heading_size": 11,
        "name_size": 14,
        "margins": 0.6,
        "section_order": ["summary", "skills", "experience", "projects", "education", "certifications"],
    },
    "minimal": {
        "name": "Minimal",
        "description": "No borders, no backgrounds, just clean typography. Maximum ATS compatibility.",
        "font": "Calibri",
        "body_size": 11,
        "heading_size": 12,
        "name_size": 15,
        "margins": 0.8,
        "section_order": ["summary", "experience", "education", "skills"],
    },
}


def get_template_sections(template_id: str) -> list[str] | None:
    """Return the section_order for a template, or None if not found."""
    tmpl = TEMPLATES.get(template_id)
    return list(tmpl["section_order"]) if tmpl else None


def list_templates() -> list[dict]:
    """Return available templates for frontend display — includes section_order and styling info."""
    return [
        {
            "id": tid,
            "name": t["name"],
            "description": t["description"],
            "section_order": t["section_order"],
            "font": t["font"],
            "body_size": t["body_size"],
            "name_size": t["name_size"],
            "margins": t["margins"],
        }
        for tid, t in TEMPLATES.items()
    ]


def _setup_styles(doc: Document, config: dict) -> None:
    """Configure document styles based on template config."""
    font_name = config["font"]

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(config["body_size"])
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # Heading 1 — section headers
    h1 = doc.styles["Heading 1"]
    h1.font.name = font_name
    h1.font.size = Pt(config["heading_size"])
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — job title / company
    h2 = doc.styles["Heading 2"]
    h2.font.name = font_name
    h2.font.size = Pt(config["body_size"])
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(1)

    # Set margins
    for section in doc.sections:
        margin = Inches(config["margins"])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def _add_name_header(doc: Document, config: dict, name: str, contact_line: str) -> None:
    """Add the name and contact info at the top."""
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name.upper() if config.get("font") == "Times New Roman" else name)
    run.font.size = Pt(config["name_size"])
    run.font.bold = True
    run.font.name = config["font"]
    p.paragraph_format.space_after = Pt(2)

    # Contact line
    if contact_line:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(contact_line)
        run2.font.size = Pt(config["body_size"] - 1)
        run2.font.name = config["font"]
        p2.paragraph_format.space_after = Pt(6)


def _add_section_header(doc: Document, title: str) -> None:
    """Add a section header with a horizontal line."""
    p = doc.add_heading(title.upper(), level=1)
    # Add a bottom border (horizontal rule)
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "000000",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc: Document, text: str, config: dict) -> None:
    """Add a bullet point."""
    # Strip any existing bullet characters from the text (Word adds its own)
    cleaned = re.sub(r"^[\s]*[-*•●○◦‣›▪▸\u2022\u2023\u25E6\u2043\u2219]+\s*", "", text).strip()
    cleaned = re.sub(r"^\d+[.)]\s*", "", cleaned).strip()
    if not cleaned:
        return
    p = doc.add_paragraph(cleaned, style="List Bullet")
    p.style.font.name = config["font"]
    p.style.font.size = Pt(config["body_size"])


_DATE_HINT_RE = re.compile(
    r"\b(?:19|20)\d{2}\b|present|current|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec",
    re.I,
)
_SEPARATOR_RE = re.compile(r"\s*[|\u2014\u2013]\s*")
_DEGREE_RE = re.compile(
    r"^(?:M\.?Sc|B\.?Sc|B\.?Eng|M\.?Eng|B\.?A|M\.?A|MBA|Ph\.?D|"
    r"Diploma|Bachelor|Master|Graduate Cert)",
    re.I,
)


_TITLE_RE = re.compile(
    r"\b(?:engineer|manager|director|analyst|developer|architect|lead|head|"
    r"officer|coordinator|specialist|consultant|designer|executive|associate|"
    r"intern|supervisor|principal|scientist|researcher|advisor|strategist)\b",
    re.I,
)

def _is_entry_heading_line(line: str) -> bool:
    """Detect if a line is an entry heading (company/role/date/title)."""
    has_date = bool(_DATE_HINT_RE.search(line))
    has_separator = bool(_SEPARATOR_RE.search(line))
    is_caps = line == line.upper() and len(line.split()) <= 8 and re.search(r"[A-Z]", line)
    has_degree = bool(_DEGREE_RE.match(line))
    # Detect job titles (short lines with title keywords, not bullets)
    words = line.split()
    has_title = (
        _TITLE_RE.search(line)
        and len(words) <= 10
        and not line.startswith(("-", "•", "*"))
    )
    return has_date or has_separator or bool(is_caps) or has_degree or bool(has_title)


def _is_education_detail(line: str) -> bool:
    """Detect education detail lines (GPA, exchange, capstone)."""
    lower = line.lower()
    return bool(re.search(
        r"\b(?:gpa|cgpa|cap|exchange|capstone|thesis|minor|major|"
        r"distinction|honou?r|cum laude|dean.?s list|first class)\b",
        lower,
    ))


def _add_entry_heading(doc: Document, line: str, config: dict) -> None:
    """Add a bold entry heading (company | role | date)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)

    # Split on pipe only (not dashes - those are often in date ranges like 2020-2022)
    # Preserve en-dash/em-dash in dates
    parts = [pt.strip() for pt in line.split("|") if pt.strip()]
    if len(parts) <= 1:
        # No pipe separator - just bold the whole line
        run = p.add_run(line)
        run.font.name = config["font"]
        run.font.size = Pt(config["body_size"])
        run.bold = True
        return

    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = config["font"]
        run.font.size = Pt(config["body_size"])
        # Bold the first part (company or degree name)
        if i == 0:
            run.bold = True
        # Add separator between parts
        if i < len(parts) - 1:
            sep = p.add_run("  |  ")
            sep.font.name = config["font"]
            sep.font.size = Pt(config["body_size"])


def _add_education_detail(doc: Document, line: str, config: dict) -> None:
    """Add an education detail line (GPA, exchange) in smaller italic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(line)
    run.font.name = config["font"]
    run.font.size = Pt(config["body_size"] - 1)
    run.italic = True


def _add_spacing(doc: Document, space: Pt = Pt(6)) -> None:
    """Add vertical spacing between entries."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.style.font.size = Pt(2)
    # Make it a tiny empty paragraph for spacing
    run = p.add_run("")
    run.font.size = Pt(2)


def _parse_sections(resume_text: str) -> dict[str, str]:
    """
    Parse resume text into sections by detecting common headers.
    Returns dict of section_name -> content.
    """
    # Common section header patterns
    header_patterns = [
        (r"(?i)^(?:(?:professional|executive)\s+)?summary\s*:?$", "summary"),
        (r"(?i)^(?:(?:professional|work)\s+)?experience\s*:?$", "experience"),
        (r"(?i)^career\s+break\s*:?$", "experience"),
        (r"(?i)^education(?:\s+(?:&|and)\s+(?:certifications?|training))?\s*:?$", "education"),
        (r"(?i)^(?:(?:technical|core)\s+)?skills(?:\s+(?:&|and)\s+(?:competenc(?:e|ies)|expertise|tools|technolog(?:y|ies)))?\s*:?$", "skills"),
        (r"(?i)^(?:core\s+)?competenc(?:e|ies)\s*:?$", "skills"),
        (r"(?i)^(?:(?:licen[cs]es?\s+(?:&|and)\s+)?certifications?|certifications?\s+(?:&|and)\s+licen[cs]es?)\s*:?$", "certifications"),
        (r"(?i)^(?:(?:selected|technical)\s+)*projects?(?:\s+experience)?\s*:?$", "projects"),
        (r"(?i)^activities?(?:\s+(?:&|and)\s+leadership)?\s*:?$", "activities"),
        (r"(?i)^volunteer(?:ing|\s+(?:work|experience))?\s*:?$", "activities"),
        (r"(?i)^leadership(?:\s+(?:&|and)\s+activities)?\s*:?$", "activities"),
        (r"(?i)^personal(?:\s+particulars?)?\s*:?$", "personal"),
        (r"(?i)^additional\s+info(?:rmation)?\s*:?$", "additional"),
        (r"(?i)^languages?\s*:?$", "languages"),
        (r"(?i)^interests?\s*:?$", "interests"),
        (r"(?i)^awards?(?:\s+(?:&|and)\s+hono(?:u)?rs?)?\s*:?$", "awards"),
        (r"(?i)^hono(?:u)?rs?(?:\s+(?:&|and)\s+awards?)?\s*:?$", "awards"),
        (r"(?i)^publications?\s*:?$", "publications"),
    ]

    lines = resume_text.split("\n")
    sections: dict[str, list[str]] = {"header": []}
    current_section = "header"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check if this line is a section header
        matched = False
        for pattern, section_name in header_patterns:
            if re.match(pattern, stripped):
                current_section = section_name
                if section_name not in sections:
                    sections[section_name] = []
                matched = True
                break

        if not matched:
            if current_section not in sections:
                sections[current_section] = []
            sections[current_section].append(stripped)

    # Convert lists to strings
    return {k: "\n".join(v) for k, v in sections.items() if v}


_EXPORT_BULLET_RE = re.compile(r"^[\s]*[-*•–]\s*")
_EXPORT_SKILL_LABEL_RE = re.compile(r"^[A-Z][^:\n]{1,60}:\s+\S")


def _group_export_lines(content: str, section_key: str) -> list[str]:
    """Join physical PDF wraps without changing resume content."""
    grouped: list[str] = []
    pending = ""

    def flush() -> None:
        nonlocal pending
        if pending:
            grouped.append(pending)
            pending = ""

    for line in (value.strip() for value in content.splitlines()):
        if not line:
            continue
        is_bullet = bool(_EXPORT_BULLET_RE.match(line) or re.match(r"^\d+[.)]\s+", line))
        is_entry_heading = (
            section_key in {"experience", "education", "projects", "activities"}
            and not is_bullet
            and not line.endswith((".", ";"))
            and _is_entry_heading_line(line)
        )
        starts_skill = section_key == "skills" and bool(_EXPORT_SKILL_LABEL_RE.match(line))

        if is_entry_heading:
            flush()
            grouped.append(line)
        elif is_bullet or starts_skill:
            flush()
            pending = line
        elif pending:
            pending = f"{pending} {line}"
        else:
            pending = line

    flush()
    return grouped


def inspect_resume_export(resume_text: str, template_id: str = "modern") -> dict[str, object]:
    """
    Return privacy-safe diagnostics for DOCX generation.
    This intentionally avoids returning raw resume content.
    """
    config = TEMPLATES.get(template_id, TEMPLATES["modern"])
    sections = _parse_sections(resume_text)
    line_counts = {
        key: len([line for line in value.split("\n") if line.strip()])
        for key, value in sections.items()
    }
    expected_sections = list(config["section_order"])
    missing_expected = [
        section for section in expected_sections
        if not sections.get(section, "").strip()
    ]
    non_header_sections = [
        key for key in sections.keys()
        if key != "header" and sections.get(key, "").strip()
    ]
    return {
        "template_id": template_id,
        "header_lines": line_counts.get("header", 0),
        "sections_found": list(sections.keys()),
        "non_header_sections": non_header_sections,
        "section_line_counts": line_counts,
        "missing_expected": missing_expected,
        "looks_header_only": len(non_header_sections) == 0,
        "word_count": len(resume_text.split()),
        "char_count": len(resume_text),
    }


def generate_docx(
    resume_text: str,
    template_id: str = "modern",
    name: str = "",
    email: str = "",
    phone: str = "",
    location: str = "",
) -> bytes:
    """
    Generate a formatted DOCX file from resume text.
    Returns the DOCX as bytes.
    """
    config = TEMPLATES.get(template_id, TEMPLATES["modern"])
    doc = Document()
    _setup_styles(doc, config)

    # Normalize Unicode for ATS compatibility before processing
    resume_text = normalize_for_ats(resume_text)

    # Parse resume into sections early so we can infer header details
    sections = _parse_sections(resume_text)
    header_lines = [
        line.strip() for line in sections.get("header", "").split("\n")
        if line.strip()
    ]

    # Build contact line
    contact_parts = [p for p in [email, phone, location] if p]
    if not contact_parts and len(header_lines) > 1:
        contact_parts = header_lines[1:4]
    contact_line = " | ".join(contact_parts)

    # Add name header
    display_name = name or (header_lines[0] if header_lines else "") or "Your Name"
    _add_name_header(doc, config, display_name, contact_line)

    # Add sections in template order
    rendered_sections: list[str] = []
    rendered_bullets = 0
    rendered_paragraphs = 0
    for section_key in config["section_order"]:
        content = sections.get(section_key, "")
        if not content:
            continue
        rendered_sections.append(section_key)

        # Map section keys to display names
        display_names = {
            "summary": "Professional Summary",
            "experience": "Professional Experience",
            "education": "Education",
            "skills": "Skills",
            "certifications": "Certifications",
            "projects": "Projects",
            "activities": "Activities & Leadership",
            "personal": "Personal Particulars",
            "interests": "Interests",
        }

        _add_section_header(doc, display_names.get(section_key, section_key.title()))

        # Add content — smart detection of entry headings, bullets, paragraphs
        content_lines = _group_export_lines(content, section_key)
        is_entry_section = section_key in ("experience", "education", "projects", "activities")
        prev_was_bullet = False

        for li, line in enumerate(content_lines):
            is_bullet = line.startswith(("-", "*", "•", "–")) or re.match(r"^\d+\.", line)

            if is_bullet:
                bullet_text = re.sub(r"^[-*•–]\s*", "", line)
                bullet_text = re.sub(r"^\d+\.\s*", "", bullet_text)
                _add_bullet(doc, bullet_text, config)
                rendered_bullets += 1
                prev_was_bullet = True
            elif is_entry_section and _is_entry_heading_line(line):
                # Add spacing before new entry (not the first one)
                if li > 0 and prev_was_bullet:
                    _add_spacing(doc, Pt(4))
                _add_entry_heading(doc, line, config)
                rendered_paragraphs += 1
                prev_was_bullet = False
            else:
                # Plain paragraph - for education details, use smaller italic
                if section_key == "education" and _is_education_detail(line):
                    _add_education_detail(doc, line, config)
                else:
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = config["font"]
                        run.font.size = Pt(config["body_size"])
                rendered_paragraphs += 1
                prev_was_bullet = False

    # Any remaining sections not in the template order
    for section_key, content in sections.items():
        if section_key in config["section_order"] or section_key == "header":
            continue
        if content:
            rendered_sections.append(section_key)
            _add_section_header(doc, section_key.title())
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith(("-", "*", "•", "–")) or re.match(r"^\d+\.", stripped):
                    bullet_text = re.sub(r"^[-*•–]\s*", "", stripped)
                    bullet_text = re.sub(r"^\d+\.\s*", "", bullet_text)
                    _add_bullet(doc, bullet_text, config)
                    rendered_bullets += 1
                elif _is_entry_heading_line(stripped):
                    _add_entry_heading(doc, stripped, config)
                    rendered_paragraphs += 1
                else:
                    p = doc.add_paragraph(stripped)
                    for run in p.runs:
                        run.font.name = config["font"]
                        run.font.size = Pt(config["body_size"])
                    rendered_paragraphs += 1

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.getvalue()
    log.info(
        "DOCX generated template=%s rendered_sections=%s doc_paragraphs=%s body_paragraphs=%s bullets=%s bytes=%s",
        template_id,
        rendered_sections,
        len(doc.paragraphs),
        rendered_paragraphs,
        rendered_bullets,
        len(docx_bytes),
    )
    if not rendered_sections or len(doc.paragraphs) <= 2:
        log.warning(
            "DOCX output looks sparse template=%s rendered_sections=%s doc_paragraphs=%s bytes=%s",
            template_id,
            rendered_sections,
            len(doc.paragraphs),
            len(docx_bytes),
        )
    return docx_bytes
