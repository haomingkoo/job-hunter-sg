from __future__ import annotations

import asyncio
from io import BytesIO
from pathlib import Path
import threading
import zipfile

import pytest
from fastapi import HTTPException, UploadFile
from starlette.datastructures import Headers


def _http_scope(*, path: str = "/api/resume/upload", headers=None) -> dict:
    return {
        "type": "http",
        "asgi": {"version": "3.0"},
        "http_version": "1.1",
        "method": "POST",
        "scheme": "http",
        "path": path,
        "raw_path": path.encode(),
        "query_string": b"",
        "headers": headers or [],
        "client": ("127.0.0.1", 50000),
        "server": ("testserver", 80),
    }


async def _run_asgi(app, scope: dict, request_messages=None) -> list[dict]:
    requests = iter(request_messages or [{"type": "http.request", "body": b""}])
    responses: list[dict] = []

    async def receive() -> dict:
        return next(requests)

    async def send(message: dict) -> None:
        responses.append(message)

    await app(scope, receive, send)
    return responses


async def _read_request_body(_scope, receive, send) -> None:
    body = bytearray()
    more_body = True
    while more_body:
        message = await receive()
        body.extend(message.get("body", b""))
        more_body = message.get("more_body", False)
    await send({"type": "http.response.start", "status": 200, "headers": []})
    await send({"type": "http.response.body", "body": bytes(body)})


def test_body_limit_rejects_oversized_content_length_without_reading_body():
    from security import RequestBodyLimitMiddleware

    app = RequestBodyLimitMiddleware(_read_request_body, default_max_bytes=5)
    responses = asyncio.run(_run_asgi(app, _http_scope(headers=[(b"content-length", b"6")])))

    assert responses[0]["status"] == 413


def test_body_limit_rejects_chunked_body_without_content_length():
    from security import RequestBodyLimitMiddleware

    app = RequestBodyLimitMiddleware(_read_request_body, default_max_bytes=5)
    responses = asyncio.run(
        _run_asgi(
            app,
            _http_scope(),
            [
                {"type": "http.request", "body": b"abc", "more_body": True},
                {"type": "http.request", "body": b"def", "more_body": False},
            ],
        )
    )

    assert responses[0]["status"] == 413


def test_resume_upload_reads_at_most_five_megabytes_plus_sentinel():
    from main import MAX_FILE_SIZE, upload_resume

    upload = UploadFile(
        file=BytesIO(b"x" * (MAX_FILE_SIZE + 1)),
        filename="resume.pdf",
        headers=Headers({"content-type": "application/pdf"}),
    )

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(upload_resume(file=upload, user=None, db=None))

    assert exc_info.value.status_code == 413
    assert upload.file.closed


def test_resume_upload_parsing_runs_outside_event_loop(monkeypatch):
    import main

    caller_thread = threading.get_ident()
    parser_threads: list[int] = []

    def fake_parse_resume(**_kwargs) -> dict:
        parser_threads.append(threading.get_ident())
        return {
            "text": "Parsed resume",
            "file_type": "pdf",
            "word_count": 2,
            "line_count": 1,
            "parse_quality": {},
        }

    class FakeSession:
        def add(self, _value) -> None:
            pass

        def commit(self) -> None:
            pass

    monkeypatch.setattr(main, "parse_resume_isolated", fake_parse_resume)
    upload = UploadFile(
        file=BytesIO(b"small-pdf"),
        filename="resume.pdf",
        headers=Headers({"content-type": "application/pdf"}),
    )

    result = asyncio.run(main.upload_resume(file=upload, user=None, db=FakeSession()))

    assert result["text"] == "Parsed resume"
    assert parser_threads and parser_threads[0] != caller_thread


def test_resume_upload_rejects_when_parser_is_busy(monkeypatch):
    import main

    monkeypatch.setattr(main, "_RESUME_PARSE_SLOTS", threading.BoundedSemaphore(0))
    upload = UploadFile(
        file=BytesIO(b"small-pdf"),
        filename="resume.pdf",
        headers=Headers({"content-type": "application/pdf"}),
    )

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(main.upload_resume(file=upload, user=None, db=None))

    assert exc_info.value.status_code == 503
    assert exc_info.value.headers == {"Retry-After": "2"}


def test_isolated_parser_returns_normal_pdf_result():
    from resume_parser import parse_resume_isolated

    path = Path(__file__).parents[1] / "templates/nus/NUS Guidelines.pdf"
    result = parse_resume_isolated(path.name, "application/pdf", path.read_bytes())

    assert result["file_type"] == "pdf"
    assert result["text"]


def test_isolated_parser_returns_normal_docx_result():
    from docx import Document
    from resume_parser import parse_resume_isolated

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com")
    document.add_paragraph("EXPERIENCE")
    document.add_paragraph("Software Engineer at Example")
    buffer = BytesIO()
    document.save(buffer)

    result = parse_resume_isolated(
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )

    assert result["file_type"] == "docx"
    assert result["name"] == "Jane Doe"


def test_docx_zip_bomb_is_rejected_before_parsing():
    from resume_parser import extract_text_from_docx

    archive_bytes = BytesIO()
    with zipfile.ZipFile(archive_bytes, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"A" * (2 * 1024 * 1024))

    with pytest.raises(ValueError, match="compression ratio is unsafe"):
        extract_text_from_docx(archive_bytes.getvalue())
