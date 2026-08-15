from __future__ import annotations

import os
import io
import secrets
import sys
from datetime import datetime, timezone
from io import BytesIO

from fastapi.testclient import TestClient
from openpyxl import Workbook

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


def _signup(client: TestClient) -> dict:
    from auth import create_token
    from database import SessionLocal
    from models import User

    with SessionLocal() as db:
        user = User(
            email=f"workspace_{secrets.token_hex(4)}@aisg.sg",
            password_hash="test-only",  # pragma: allowlist secret
            name="Workspace User",
            email_verified_at=datetime.now(timezone.utc),
        )
        db.add(user)
        db.commit()
        db.refresh(user)
        token = create_token(user.id, user.token_version)
    return {"Authorization": f"Bearer {token}"}


def _create_workspace_with_resume(client: TestClient, headers: dict, role_metadata: dict | None = None) -> dict:
    resume = client.post("/api/resume/versions", json={
        "label": "Master resume",
        "resume_text": "EXPERIENCE\n- Built Python data pipelines for public-sector services.",
        "is_master": True,
    }, headers=headers)
    assert resume.status_code == 200

    created = client.post("/api/applications/workspaces", json={
        "company": "GovTech",
        "title": "Senior AI Engineer",
        "job_description": "Build agentic workflows for public-sector digital services.",
        "status": "saved",
        "resume_version_id": resume.json()["id"],
        "role_metadata": role_metadata or {},
    }, headers=headers)
    assert created.status_code == 201
    return created.json()


def _mom_workbook(occupation: str = "Artificial intelligence engineer") -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "T4"
    sheet.cell(row=10, column=3, value=occupation)
    for column, value in enumerate((6000, 8000, 10000, 6200, 8500, 10500), start=4):
        sheet.cell(row=10, column=column, value=value)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def test_application_workspace_stores_job_context_and_append_only_history():
    from database import init_db
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)

    version = client.post("/api/resume/versions", json={
        "label": "Master resume",
        "resume_text": "Built AI job-search workflows with Python, React, FastAPI, and reliable tests.",
        "is_master": True,
    }, headers=headers)
    assert version.status_code == 200
    stored_version = client.get(
        f"/api/resume/versions/{version.json()['id']}",
        headers=headers,
    ).json()
    from resume_document import is_resume_document

    assert is_resume_document(stored_version["resume_structured"])
    assert stored_version["resume_structured"]["raw_text"] == stored_version["resume_text"]

    created = client.post("/api/applications/workspaces", json={
        "company": "GovTech",
        "title": "Senior AI Engineer",
        "job_description": "Build agentic workflows for public-sector digital services.",
        "source_url": "https://example.com/jobs/ai-engineer",
        "source": "manual",
        "status": "saved",
        "date_applied": "2026-07-04",
        "resume_version_id": version.json()["id"],
        "role_metadata": {"seniority": "senior", "location": "Singapore"},
    }, headers=headers)
    assert created.status_code == 201
    workspace = created.json()

    assert workspace["company"] == "GovTech"
    assert workspace["title"] == "Senior AI Engineer"
    assert workspace["job_description"] == "Build agentic workflows for public-sector digital services."
    assert workspace["source_url"] == "https://example.com/jobs/ai-engineer"
    assert workspace["resume_version_id"] == version.json()["id"]
    assert workspace["role_metadata"] == {"seniority": "senior", "location": "Singapore"}
    assert [event["stage"] for event in workspace["stage_history"]] == ["saved"]

    updated = client.put(
        f"/api/tracked/{workspace['id']}",
        json={"status": "interview"},
        headers=headers,
    )
    assert updated.status_code == 200

    loaded = client.get(f"/api/applications/workspaces/{workspace['id']}", headers=headers)
    assert loaded.status_code == 200
    assert [event["stage"] for event in loaded.json()["stage_history"]] == [
        "saved",
        "interview",
    ]


def test_cover_letter_persists_exact_resume_provenance_and_latest_candidate_edits(monkeypatch):
    from database import init_db
    import main

    init_db()
    client = TestClient(main.app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)
    prompts = []
    monkeypatch.setattr(main, "_consume_ai_credit", lambda *args, **kwargs: None)

    def draft(messages, **_kwargs):
        prompts.append(messages[1]["content"])
        return "Dear Hiring Team, " + "My Python delivery experience aligns with this public-sector role. " * 8

    monkeypatch.setattr(main, "_call_sealion", draft)
    generated = client.post("/api/ai/cover-letter", json={
        "workspace_id": workspace["id"],
        "job_title": workspace["title"],
        "job_company": workspace["company"],
    }, headers=headers)

    assert generated.status_code == 200
    assert generated.json()["saved"] is True
    assert generated.json()["resume_version_id"] == workspace["resume_version_id"]
    assert "Built Python data pipelines" in prompts[0]

    loaded = client.get(f"/api/applications/workspaces/{workspace['id']}", headers=headers)
    document = loaded.json()["role_metadata"]["cover_letter"]
    assert document["resume_version_id"] == workspace["resume_version_id"]
    assert document["content"] == generated.json()["cover_letter"]

    edited_content = "Dear Hiring Team,\n\n" + "This candidate-verified edit remains grounded in my stored experience. " * 8
    edited = client.put(
        f"/api/applications/workspaces/{workspace['id']}/cover-letter",
        json={"content": edited_content},
        headers=headers,
    )
    assert edited.status_code == 200
    assert edited.json()["content"] == edited_content.strip()
    assert edited.json()["generated_at"] == document["generated_at"]

    replaced = client.post("/api/ai/cover-letter", json={
        "resume_text": "UNTRUSTED FALLBACK " * 10,
        "workspace_id": workspace["id"],
        "job_title": workspace["title"],
        "job_company": workspace["company"],
    }, headers=headers)
    assert replaced.status_code == 200
    reloaded = client.get(f"/api/applications/workspaces/{workspace['id']}", headers=headers).json()
    assert isinstance(reloaded["role_metadata"]["cover_letter"], dict)
    assert reloaded["role_metadata"]["cover_letter"]["content"] == replaced.json()["cover_letter"]

    unlinked = client.post("/api/applications/workspaces", json={
        "company": "Unlinked Company",
        "title": "AI Platform Engineer",
        "job_description": "Build reliable production AI platforms.",
    }, headers=headers)
    assert unlinked.status_code == 201
    unlinked_generation = client.post("/api/ai/cover-letter", json={
        "resume_text": "Built reliable Python and AI systems for production teams. " * 3,
        "workspace_id": unlinked.json()["id"],
        "job_title": "AI Platform Engineer",
        "job_company": "Unlinked Company",
    }, headers=headers)
    assert unlinked_generation.status_code == 200
    assert unlinked_generation.json()["resume_version_id"] is None
    unlinked_loaded = client.get(
        f"/api/applications/workspaces/{unlinked.json()['id']}", headers=headers,
    ).json()
    assert unlinked_loaded["role_metadata"]["cover_letter"]["resume_version_id"] is None

    other_headers = _signup(client)
    denied_generation = client.post("/api/ai/cover-letter", json={
        "resume_text": "Built reliable Python and AI systems for production teams. " * 3,
        "workspace_id": workspace["id"],
        "job_title": workspace["title"],
        "job_company": workspace["company"],
    }, headers=other_headers)
    assert denied_generation.status_code == 404
    denied = client.put(
        f"/api/applications/workspaces/{workspace['id']}/cover-letter",
        json={"content": edited_content},
        headers=other_headers,
    )
    assert denied.status_code == 404


def test_cover_letter_rejects_mismatch_and_failed_generation_stores_nothing(monkeypatch):
    from database import SessionLocal, init_db
    from models import ScrapedJob
    import main

    init_db()
    client = TestClient(main.app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)
    monkeypatch.setattr(main, "_consume_ai_credit", lambda *args, **kwargs: None)
    monkeypatch.setattr(main, "_call_sealion", lambda *args, **kwargs: None)

    mismatch = client.post("/api/ai/cover-letter", json={
        "resume_text": "Fallback resume content with enough detail for request validation. " * 2,
        "workspace_id": workspace["id"],
        "job_id": 999999,
        "job_title": workspace["title"],
        "job_company": workspace["company"],
    }, headers=headers)
    assert mismatch.status_code == 409

    marker = secrets.token_hex(6)
    with SessionLocal() as db:
        original = ScrapedJob(
            title="AI Engineer",
            company="Same Company",
            description="Original posting",
            dedup_key=f"cover-original-{marker}",
        )
        repost = ScrapedJob(
            title="AI Engineer",
            company="Same Company",
            description="Reposted role",
            dedup_key=f"cover-repost-{marker}",
        )
        db.add_all([original, repost])
        db.commit()
        original_id, repost_id = original.id, repost.id

    repost_workspace = client.post("/api/applications/workspaces", json={
        "company": "Same Company",
        "title": "AI Engineer",
        "job_description": "Original posting",
        "scraped_job_id": original_id,
        "resume_version_id": workspace["resume_version_id"],
    }, headers=headers)
    assert repost_workspace.status_code == 201
    repost_collision = client.post("/api/ai/cover-letter", json={
        "resume_text": "",
        "workspace_id": repost_workspace.json()["id"],
        "job_id": repost_id,
        "job_title": "AI Engineer",
        "job_company": "Same Company",
    }, headers=headers)
    assert repost_collision.status_code == 409

    failed = client.post("/api/ai/cover-letter", json={
        "resume_text": "Fallback resume content with enough detail for request validation. " * 2,
        "workspace_id": workspace["id"],
        "job_title": workspace["title"],
        "job_company": workspace["company"],
    }, headers=headers)
    assert failed.status_code == 503
    loaded = client.get(f"/api/applications/workspaces/{workspace['id']}", headers=headers)
    assert "cover_letter" not in loaded.json()["role_metadata"]


def test_application_outcomes_are_queryable_by_status_history_and_resume_version():
    from database import init_db
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)
    version_id = workspace["resume_version_id"]

    moved = client.put(
        f"/api/tracked/{workspace['id']}",
        json={"status": "interview"},
        headers=headers,
    )
    assert moved.status_code == 200

    for company, status, resume_version_id in [
        ("Submitted Co", "applied", version_id),
        ("Offer Co", "accepted", version_id),
        ("Rejected Co", "rejected", None),
        ("Withdrawn Co", "withdrawn", None),
        ("Silent Co", "no_response", None),
    ]:
        response = client.post("/api/tracked", json={
            "company": company,
            "role": "AI Program Manager",
            "date_applied": "2026-07-04",
            "status": status,
            "resume_version_id": resume_version_id,
        }, headers=headers)
        assert response.status_code == 201

    response = client.get("/api/applications/outcomes", headers=headers)
    assert response.status_code == 200
    summary = response.json()

    assert summary["total_applications"] == 6
    assert summary["counts"] == {
        "submitted": 1,
        "interview": 1,
        "offer": 1,
        "rejected": 1,
        "withdrawn": 1,
        "no_response": 1,
    }
    assert summary["stage_counts"]["interview"] == 1
    assert summary["stage_counts"]["submitted"] == 1
    assert summary["unlinked_applications"] == 3
    assert summary["resume_versions"] == [
        {
            "resume_version_id": version_id,
            "applications": 3,
            "counts": {
                "submitted": 1,
                "interview": 1,
                "offer": 1,
                "rejected": 0,
                "withdrawn": 0,
                "no_response": 0,
            },
        }
    ]


def test_application_workspace_requires_company_title_and_job_description():
    from database import init_db
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)

    response = client.post("/api/applications/workspaces", json={
        "company": "GovTech",
        "title": "Senior AI Engineer",
    }, headers=headers)

    assert response.status_code == 422
    assert "job_description" in response.text


def test_application_workspace_module_creates_and_moves_workspace():
    from application_workspace import (
        create_application_workspace,
        get_application_workspace,
        update_tracked_job,
    )
    from database import SessionLocal, init_db
    from models import User
    from schemas import ApplicationWorkspaceCreate, TrackedJobUpdate

    init_db()
    db = SessionLocal()
    try:
        from auth import hash_password

        user = User(
            email=f"workspace_module_{secrets.token_hex(4)}@aisg.sg",
            password_hash=hash_password("TestPassword123!"),
            name="Workspace Module User",
            tier="pro",
        )
        db.add(user)
        db.commit()
        db.refresh(user)

        workspace = create_application_workspace(
            db,
            user,
            ApplicationWorkspaceCreate(
                company="GovTech",
                title="Senior AI Engineer",
                job_description="Build agentic workflows for public-sector digital services.",
                status="saved",
            ),
        )
        assert workspace["title"] == "Senior AI Engineer"
        assert [event["stage"] for event in workspace["stage_history"]] == ["saved"]

        moved = update_tracked_job(
            db,
            user,
            workspace["id"],
            TrackedJobUpdate(status="interview"),
        )
        loaded = get_application_workspace(db, user.id, moved.id)

        assert loaded["status"] == "interview"
        assert [event["stage"] for event in loaded["stage_history"]] == [
            "saved",
            "interview",
        ]
    finally:
        db.close()


def test_recruitment_pipeline_reuses_record_and_preserves_user_workspace_data():
    from application_workspace import create_tracked_job, ensure_recruitment_application
    from database import SessionLocal, init_db
    from models import ResumeVersion, TrackedJob, User
    from schemas import TrackedJobCreate

    init_db()
    db = SessionLocal()
    try:
        user = User(
            email=f"pipeline_{secrets.token_hex(4)}@aisg.sg",
            password_hash="test-only",  # pragma: allowlist secret
            name="Pipeline User",
            tier="pro",
        )
        db.add(user)
        db.flush()
        resume = ResumeVersion(
            user_id=user.id,
            label="Operations resume",
            resume_text="Led semiconductor manufacturing transformation.",
            is_master=True,
        )
        db.add(resume)
        db.commit()

        existing = create_tracked_job(
            db,
            user,
            TrackedJobCreate(
                company="Example Semiconductor",
                role="Operations Manager",
                status="interview",
                notes="Follow up with the hiring manager.",
                role_metadata={
                    "contacts": [{"name": "Hiring Manager", "details": "Introduced at SEMICON"}],
                    "activity": [{"type": "contact_added", "recorded_at": "2026-08-01T00:00:00Z"}],
                },
                resume_version_id=resume.id,
            ),
        )
        original_history = list(existing.stage_history)

        reused = ensure_recruitment_application(
            db,
            user,
            TrackedJobCreate(
                company="Example Semiconductor",
                role="Operations Manager",
                status="saved",
                source="MyCareersFuture",
                source_url="https://example.test/jobs/901",
                job_description="Lead fab operations and continuous improvement.",
                scraped_job_id=901,
                resume_version_id=resume.id,
            ),
            thread_id="thread-pipeline",
            source_job_id=901,
            posting_snapshot={"job_id": 901, "description": "Original posting snapshot"},
            fit_evidence={"matched": [{"statement": "Manufacturing leadership"}]},
            selected=True,
            existing_tracked_job_id=existing.id,
        )
        db.commit()
        db.refresh(reused)

        assert db.query(TrackedJob).filter(TrackedJob.user_id == user.id).count() == 1
        assert reused.id == existing.id
        assert reused.status == "interview"
        assert reused.notes == "Follow up with the hiring manager."
        assert reused.stage_history == original_history
        assert reused.role_metadata["contacts"][0]["name"] == "Hiring Manager"
        assert reused.role_metadata["activity"][0]["type"] == "contact_added"
        assert reused.role_metadata["recruitment_pipeline"]["state"] == "selected"
        assert reused.role_metadata["recruitment_pipeline"]["posting_snapshot"]["job_id"] == 901
    finally:
        db.close()


def test_application_workspace_agent_review_saves_artifacts(monkeypatch):
    from database import init_db
    import main
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(
        client,
        headers,
        role_metadata={"submitted_resume": {"artifact_id": "submitted-1"}},
    )
    workspace_id = workspace["id"]
    source_resume_version_id = workspace["resume_version_id"]

    seen_body = {}

    def fake_stream(body):
        seen_body.update(body)
        yield {"event": "session", "session_id": "workspace-agent-review"}
        yield {"event": "token", "session_id": "workspace-agent-review", "content": "Emphasize agentic workflow delivery."}
        yield {"event": "done", "session_id": "workspace-agent-review"}

    monkeypatch.setattr(main, "_stream_resume_agent_events", fake_stream)
    monkeypatch.setattr(
        main,
        "_get_resume_agent_state",
        lambda session_id, owner_key=None: {
            "session_id": session_id,
            "pending_diffs": [
                {
                    "bullet_id": "exp-0-b0",
                    "original": "Built Python data pipelines for public-sector services.",
                    "rewrite": "Built agentic Python data workflows for public-sector services.",
                    "status": "pending",
                }
            ],
            "debate_summary": {
                "roles": ["recruiter", "ats", "skeptic"],
                "key_disagreements": ["ATS wants more keyword coverage; skeptic wants proof before adding claims."],
                "final_recommendation": "Revise one bullet, then rerun review.",
                "confidence": "medium",
                "trace_id": "trace-123",
            },
        },
    )

    response = client.post(f"/api/applications/workspaces/{workspace_id}/agent-review", headers=headers)
    assert response.status_code == 200
    workspace = response.json()
    review = workspace["role_metadata"]["agent_review"]

    assert "GovTech" in seen_body["message"]
    assert seen_body["resume_text"] == "EXPERIENCE\n- Built Python data pipelines for public-sector services."
    assert review["status"] == "completed"
    assert review["role_brief"]["title"] == "Senior AI Engineer"
    assert review["recommendations"] == ["Emphasize agentic workflow delivery."]
    assert review["pending_diffs"][0]["bullet_id"] == "exp-0-b0"
    assert review["debate_summary"] == {
        "roles": ["recruiter", "ats", "skeptic"],
        "key_disagreements": ["ATS wants more keyword coverage; skeptic wants proof before adding claims."],
        "final_recommendation": "Revise one bullet, then rerun review.",
        "confidence": "medium",
        "trace_id": "trace-123",
    }
    assert review["tailored_draft"]["source_resume_version_id"] == source_resume_version_id
    assert workspace["role_metadata"]["submitted_resume"] == {"artifact_id": "submitted-1"}
    assert workspace["stage_history"][-1]["source"] == "agent_review"

    draft = client.get(f"/api/resume/versions/{review['tailored_draft']['resume_version_id']}", headers=headers)
    assert draft.status_code == 200
    assert draft.json()["source"] == "agent_review"
    assert "Built agentic Python data workflows for public-sector services." in draft.json()["resume_text"]

    source = client.get(f"/api/resume/versions/{source_resume_version_id}", headers=headers)
    assert source.status_code == 200
    assert "Built Python data pipelines for public-sector services." in source.json()["resume_text"]


def test_application_workspace_agent_review_returns_clean_config_error(monkeypatch):
    import resume_agent.models as agent_models
    from database import init_db
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    monkeypatch.setattr(agent_models.ai_service, "_get_api_key", lambda: "")
    workspace_id = _create_workspace_with_resume(client, headers)["id"]

    response = client.post(f"/api/applications/workspaces/{workspace_id}/agent-review", headers=headers)
    assert response.status_code == 503
    assert response.json()["detail"] == "Agent v2 needs SEALION_API_KEYS or SEALION_API configured before it can run."

    loaded = client.get(f"/api/applications/workspaces/{workspace_id}", headers=headers)
    review = loaded.json()["role_metadata"]["agent_review"]
    assert review["status"] == "error"
    assert loaded.json()["stage_history"][-1]["source"] == "agent_review_error"


def test_application_workspace_submitted_resume_artifact_survives_draft_generation(monkeypatch):
    from docx import Document

    from database import init_db
    import main
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)
    workspace_id = workspace["id"]

    doc = Document()
    doc.add_paragraph("Submitted resume for GovTech Senior AI Engineer.")
    doc.add_paragraph("Built agentic workflows with Python and FastAPI.")
    buffer = io.BytesIO()
    doc.save(buffer)

    uploaded = client.post(
        f"/api/applications/workspaces/{workspace_id}/submitted-resume",
        headers=headers,
        data={"submitted_date": "2026-07-04", "notes": "Submitted through company portal."},
        files={
            "file": (
                "submitted.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 200
    submitted = uploaded.json()["role_metadata"]["submitted_resume"]
    assert submitted["filename"] == "submitted.docx"
    assert submitted["submitted_date"] == "2026-07-04"
    assert submitted["notes"] == "Submitted through company portal."
    assert "Submitted resume for GovTech" in submitted["text"]
    assert len(uploaded.json()["role_metadata"]["submitted_resume_artifacts"]) == 1
    assert uploaded.json()["role_metadata"]["submitted_resume_artifacts"][0]["content_base64"]

    monkeypatch.setattr(
        main,
        "_stream_resume_agent_events",
        lambda _body: iter(
            [
                {"event": "session", "session_id": "submitted-artifact-review"},
                {"event": "token", "session_id": "submitted-artifact-review", "content": "Draft saved."},
                {"event": "done", "session_id": "submitted-artifact-review"},
            ]
        ),
    )
    monkeypatch.setattr(
        main,
        "_get_resume_agent_state",
        lambda session_id, owner_key=None: {
            "session_id": session_id,
            "pending_diffs": [
                {
                    "bullet_id": "exp-0-b0",
                    "original": "Built Python data pipelines for public-sector services.",
                    "rewrite": "Built agentic Python data workflows for public-sector services.",
                    "status": "pending",
                }
            ],
        },
    )

    reviewed = client.post(f"/api/applications/workspaces/{workspace_id}/agent-review", headers=headers)
    assert reviewed.status_code == 200
    metadata = reviewed.json()["role_metadata"]
    assert metadata["submitted_resume"]["artifact_id"] == submitted["artifact_id"]
    assert len(metadata["submitted_resume_artifacts"]) == 1
    assert metadata["agent_review"]["tailored_draft"]["resume_version_id"]


def test_application_workspace_rejects_submitted_resume_over_five_megabytes():
    from database import init_db
    from main import app
    from resume_parser import MAX_FILE_SIZE

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    workspace_id = _create_workspace_with_resume(client, headers)["id"]

    response = client.post(
        f"/api/applications/workspaces/{workspace_id}/submitted-resume",
        headers=headers,
        files={"file": ("oversized.pdf", b"x" * (MAX_FILE_SIZE + 1), "application/pdf")},
    )

    assert response.status_code == 413


def test_research_pack_route_persists_through_the_workspace_interface(monkeypatch):
    from application_research import ResearchPack
    from database import init_db
    from recruitment_team.http_routes import get_recruitment_telemetry
    from recruitment_team.telemetry import RecordedTelemetry
    import main

    init_db()
    client = TestClient(main.app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(
        client,
        headers,
        role_metadata={"contacts": [{"name": "Hiring manager"}]},
    )

    pack = ResearchPack(
        status="complete",
        role_company_brief={
            "status": "complete",
            "company": {"name": "GovTech"},
            "role": {"ats_terms": [{"term": "python"}]},
            "sources": [{"url": "https://example.com/job", "source_type": "job_posting"}],
        },
        interview_pack={
            "status": "complete",
            "questions": [{"cluster": "technical", "question": "How did you use Python?"}],
        },
        compensation_brief={
            "status": "complete",
            "observations": [{"kind": "employer_posting", "value": "$8,000"}],
        },
        source_statuses=({"source": "fixture", "status": "complete"},),
        built_at="2026-08-03T00:00:00+00:00",
    )

    class FakeProvider:
        def build(self, tracked, resume_text):
            assert tracked.id == workspace["id"]
            assert "Python data pipelines" in resume_text
            return pack

    telemetry = RecordedTelemetry()
    main.app.dependency_overrides[get_recruitment_telemetry] = lambda: telemetry
    monkeypatch.setattr(main, "CorpusAndMomResearchProvider", lambda _db: FakeProvider())
    built = client.post(
        f"/api/applications/workspaces/{workspace['id']}/research-pack",
        headers=headers,
    )
    assert built.status_code == 200
    metadata = built.json()["role_metadata"]
    assert metadata["contacts"] == [{"name": "Hiring manager"}]
    assert metadata["application_research"]["status"] == "complete"
    assert built.json()["stage_history"][-1]["source"] == "application_research"
    assert telemetry.spans[0].attributes == {
        "workspace_id": workspace["id"],
        "provider": "public_job_corpus_and_mom",
        "status": "complete",
        "source_count": 1,
    }

    reloaded = client.get(
        f"/api/applications/workspaces/{workspace['id']}",
        headers=headers,
    )
    assert reloaded.json()["role_metadata"]["application_research"]["built_at"] == (
        "2026-08-03T00:00:00+00:00"
    )

    other_headers = _signup(client)
    denied = client.post(
        f"/api/applications/workspaces/{workspace['id']}/research-pack",
        headers=other_headers,
    )
    assert denied.status_code == 404
    main.app.dependency_overrides.pop(get_recruitment_telemetry, None)


def test_corpus_and_mom_provider_keeps_sources_and_wage_definitions_separate():
    from application_research import CorpusAndMomResearchProvider
    from database import SessionLocal, init_db
    from models import ScrapedJob, TrackedJob

    init_db()
    unique = secrets.token_hex(6)
    with SessionLocal() as db:
        target = ScrapedJob(
            title="Artificial Intelligence Engineer",
            company=f"Evidence Company {unique}",
            location="Singapore",
            salary="$8,000 - $10,000",
            source="MyCareersFuture",
            url=f"https://example.com/{unique}/target",
            posted_date="2026-08-01",
            posted_at_sort="2026-08-01T00:00:00+00:00",
            scraped_at="2026-08-02T00:00:00+00:00",
            description="Required Python, machine learning, and stakeholder communication.",
            skills=["Python", "Machine Learning"],
            parsed_jd={"required_skills": ["Python", "Machine Learning"]},
            dedup_key=f"research-target-{unique}",
            hidden=0,
        )
        comparable = ScrapedJob(
            title="Machine Learning Engineer",
            company=f"Comparable Company {unique}",
            location="Singapore",
            salary="$7,000 - $9,000",
            source="MyCareersFuture",
            url=f"https://example.com/{unique}/comparable",
            posted_date="2026-08-01",
            posted_at_sort="2026-08-01T00:00:00+00:00",
            scraped_at="2026-08-02T00:00:00+00:00",
            description="Build Python machine learning systems with product stakeholders.",
            skills=["Python", "Machine Learning"],
            dedup_key=f"research-comparable-{unique}",
            hidden=0,
        )
        db.add_all([target, comparable])
        db.flush()
        tracked = TrackedJob(
            user_id=1,
            company=target.company,
            role=target.title,
            status="saved",
            source=target.source,
            source_url=target.url,
            job_description=target.description,
            scraped_job_id=target.id,
            role_metadata={
                "recruitment_pipeline": {
                    "posting_snapshot": {
                        "salary": target.salary,
                        "source": {"url": target.url, "posted_date": target.posted_date},
                    }
                }
            },
        )

        class Response:
            content = _mom_workbook()

            @staticmethod
            def raise_for_status():
                return None

        pack = CorpusAndMomResearchProvider(db, http_get=lambda *_args, **_kwargs: Response()).build(
            tracked,
            "Built Python machine learning systems with product and operations teams.",
        )

    assert pack.status == "complete"
    assert pack.role_company_brief["role"]["comparable_titles"]
    python_term = next(
        item for item in pack.role_company_brief["role"]["ats_terms"] if item["term"] == "python"
    )
    assert python_term["sources"]
    technical = next(
        item for item in pack.interview_pack["questions"] if item["cluster"] == "technical"
    )
    assert technical["answer_scaffold"]["evidence_quote"]
    assert technical["sources"][0]["retrieved_at"] == target.scraped_at
    assert technical["sources"][0]["evidence_note"]
    assert pack.interview_pack["answer_formats"] == ["STAR", "XYZ"]
    assert pack.interview_pack["source_state"] == "fresh"
    observations = pack.compensation_brief["observations"]
    assert [item["kind"] for item in observations] == [
        "employer_posting",
        "mom_occupational_wages",
    ]
    mom = observations[1]
    assert mom["basic_wage"] == {"p25": 6000, "median": 8000, "p75": 10000}
    assert mom["gross_wage"] == {"p25": 6200, "median": 8500, "p75": 10500}
    assert mom["excludes"] == ["bonuses", "stock options", "employer CPF"]
    assert pack.compensation_brief["comparison_state"] == "multiple_incompatible_observations"
    assert "never silently averaged" in pack.compensation_brief["comparison_rule"]
    assert [lead["publisher"] for lead in pack.compensation_brief["recruiter_guide_leads"]] == [
        "Hays Singapore",
        "Michael Page Singapore",
    ]
    assert all(lead["retrieved_at"] for lead in pack.compensation_brief["recruiter_guide_leads"])


def test_research_provider_uses_the_selected_corpus_posting_salary_without_pipeline_metadata():
    from application_research import CorpusAndMomResearchProvider
    from database import SessionLocal, init_db
    from models import ScrapedJob, TrackedJob

    init_db()
    unique = secrets.token_hex(6)
    with SessionLocal() as db:
        target = ScrapedJob(
            title="Manufacturing Manager",
            company=f"Selected Posting Company {unique}",
            salary="$8,000 - $10,000",
            source="MyCareersFuture",
            url=f"https://example.com/{unique}/selected",
            posted_date="2026-08-03",
            posted_at_sort="2026-08-03T00:00:00+00:00",
            scraped_at="2026-08-03T00:00:00+00:00",
            description="Lead manufacturing operations and continuous improvement.",
            dedup_key=f"selected-posting-pay-{unique}",
            hidden=0,
        )
        db.add(target)
        db.flush()
        tracked = TrackedJob(
            user_id=1,
            company=target.company,
            role=target.title,
            status="applied",
            source=target.source,
            source_url=target.url,
            scraped_job_id=target.id,
            role_metadata={},
        )

        class Response:
            content = _mom_workbook()

            @staticmethod
            def raise_for_status():
                return None

        pack = CorpusAndMomResearchProvider(db, http_get=lambda *_args, **_kwargs: Response()).build(tracked, "")

    posting = pack.compensation_brief["observations"][0]
    assert posting == {
        "kind": "employer_posting",
        "value": "$8,000 - $10,000",
        "currency": "SGD",
        "period": "as stated by employer",
        "definition": "Employer-stated posting range; package components were not inferred.",
        "source_url": target.url,
        "source_type": "job_posting",
        "data_date": "2026-08-03",
    }


def test_mom_role_mapping_rejects_a_shared_generic_engineer_title():
    from application_research import _mom_observation

    assert _mom_observation("System Engineer", _mom_workbook("Lift engineer")) is None


def test_research_provider_distinguishes_stale_and_sparse_niche_evidence():
    from application_research import CorpusAndMomResearchProvider
    from database import SessionLocal, init_db
    from models import ScrapedJob, TrackedJob

    init_db()
    unique = secrets.token_hex(6)
    with SessionLocal() as db:
        stale_job = ScrapedJob(
            title=f"Cryogenic Lithography Orchestrator {unique}",
            company=f"Stale Evidence Company {unique}",
            source="MyCareersFuture",
            url=f"https://example.com/{unique}/stale",
            posted_date="2024-01-01",
            posted_at_sort="2024-01-01T00:00:00+00:00",
            scraped_at="2024-01-02T00:00:00+00:00",
            description="Operate cryogenic lithography systems.",
            dedup_key=f"stale-research-{unique}",
            hidden=0,
        )
        db.add(stale_job)
        db.flush()
        tracked = TrackedJob(
            user_id=1,
            company=stale_job.company,
            role=stale_job.title,
            status="saved",
            source=stale_job.source,
            source_url=stale_job.url,
            job_description=stale_job.description,
            scraped_job_id=stale_job.id,
            role_metadata={},
        )

        class UnmatchedResponse:
            content = _mom_workbook("Registered nurse")

            @staticmethod
            def raise_for_status():
                return None

        stale_pack = CorpusAndMomResearchProvider(
            db,
            http_get=lambda *_args, **_kwargs: UnmatchedResponse(),
        ).build(tracked, "")
        niche = TrackedJob(
            user_id=1,
            company="Unknown Niche Employer",
            role="Quantum Basket Strategist",
            status="saved",
            role_metadata={},
        )
        sparse_pack = CorpusAndMomResearchProvider(
            db,
            http_get=lambda *_args, **_kwargs: UnmatchedResponse(),
        ).build(niche, "")

    assert stale_pack.role_company_brief["freshness"] == "stale"
    assert stale_pack.interview_pack["source_state"] == "stale"
    assert sparse_pack.status == "valid_empty"
    assert sparse_pack.interview_pack["status"] == "sparse"
    assert sparse_pack.interview_pack["source_state"] == "valid_empty"
    assert sparse_pack.compensation_brief["comparison_state"] == "valid_empty"
    mom_status = next(
        item for item in sparse_pack.source_statuses if item["source"] == "mom_occupational_wages_2025"
    )
    assert mom_status["status"] == "valid_empty"
    assert "No sufficiently similar" in mom_status["detail"]


def test_research_provider_exposes_access_failure_without_fabricating_results():
    import requests

    from application_research import CorpusAndMomResearchProvider
    from database import SessionLocal, init_db
    from models import TrackedJob

    init_db()
    tracked = TrackedJob(
        user_id=1,
        company="Unknown Employer",
        role="Niche Quantum Basket Strategist",
        status="saved",
        source="",
        source_url="",
        job_description="",
        role_metadata={},
    )

    def unavailable(*_args, **_kwargs):
        raise requests.ConnectionError("official source unavailable")

    with SessionLocal() as db:
        pack = CorpusAndMomResearchProvider(db, http_get=unavailable).build(tracked, "")

    assert pack.status == "access_failure"
    assert pack.role_company_brief["status"] == "valid_empty"
    assert pack.compensation_brief["observations"] == []
    statuses = {item["source"]: item["status"] for item in pack.source_statuses}
    assert statuses["public_job_corpus"] == "valid_empty"
    assert statuses["mom_occupational_wages_2025"] == "access_failure"
    assert statuses["community_and_employer_reviews"] == "valid_empty"


def test_private_negotiation_rehearsal_is_repeatable_and_never_invents_walk_away(monkeypatch):
    from database import init_db
    from main import app
    from recruitment_team.http_routes import get_recruitment_telemetry
    from recruitment_team.telemetry import RecordedTelemetry

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)
    path = f"/api/applications/workspaces/{workspace['id']}/negotiation/rehearse"
    telemetry = RecordedTelemetry()
    app.dependency_overrides[get_recruitment_telemetry] = lambda: telemetry

    def fake_coach(context):
        assert context["role"] == "Senior AI Engineer"
        assert "walk_away_point" not in context
        assert all("source_url" not in anchor for anchor in context["anchor_options"])
        assert all("data_date" not in anchor for anchor in context["anchor_options"])
        return {
            "opening": f"Respond directly to: {context['scenario']}",
            "questions": ["Which package definition applies?"],
            "trade_offs": ["Protect the stated priorities before trading another term."],
            "concessions": ["Offer flexibility only in exchange for a confirmed return."],
        }

    monkeypatch.setattr("main.coach_negotiation", fake_coach)

    first = client.post(
        path,
        headers=headers,
        json={
            "priorities": ["Role scope", "Base salary"],
            "walk_away_point": "Private candidate threshold",
            "scenario": "The recruiter says the base is fixed but bonus may move.",
            "authorized_evidence": [
                {
                    "label": "Written offer",
                    "value": "$9,000 monthly base",
                    "definition": "Monthly base excluding bonus",
                    "source_url": "https://example.com/authorized-offer",
                    "data_date": "2026-08-02",
                }
            ],
        },
    )
    assert first.status_code == 200
    negotiation = first.json()["role_metadata"]["negotiation"]
    assert negotiation["walk_away_point"] == "Private candidate threshold"
    assert negotiation["authorized_evidence"][0]["source_type"] == "self_reported_user_supplied"
    assert negotiation["authorized_evidence"][0]["data_date"] == "2026-08-02"
    assert negotiation["authorized_evidence"][0]["provided_at"]
    assert negotiation["rounds"][0]["coach_response"]["anchor_options"][0]["value"] == (
        "$9,000 monthly base"
    )
    assert "will not replace or infer" in negotiation["rounds"][0]["coach_response"][
        "walk_away_guidance"
    ]
    assert negotiation["rounds"][0]["coach_response"]["questions"]
    assert negotiation["rounds"][0]["coach_response"]["trade_offs"]
    assert negotiation["rounds"][0]["coach_response"]["concessions"]
    assert "base is fixed" in negotiation["rounds"][0]["coach_response"]["opening"]

    second = client.post(
        path,
        headers=headers,
        json={
            "priorities": ["Flexibility"],
            "scenario": "The recruiter asks for my minimum.",
        },
    )
    assert second.status_code == 200
    negotiation = second.json()["role_metadata"]["negotiation"]
    assert len(negotiation["rounds"]) == 2
    assert negotiation["walk_away_point"] == ""
    assert negotiation["rounds"][-1]["coach_response"]["walk_away_guidance"] == (
        "No walk-away point was supplied, so none was invented."
    )
    telemetry_payload = str([span.attributes for span in telemetry.spans])
    assert "Private candidate threshold" not in telemetry_payload
    assert "Written offer" not in telemetry_payload
    assert "The recruiter says" not in telemetry_payload
    assert [span.name for span in telemetry.spans] == [
        "application.negotiation_rehearsal",
        "application.negotiation_rehearsal",
    ]

    other_headers = _signup(client)
    denied = client.post(
        path,
        headers=other_headers,
        json={
            "priorities": ["Role scope"],
            "scenario": "Try to read another user's private negotiation.",
        },
    )
    assert denied.status_code == 404
    app.dependency_overrides.pop(get_recruitment_telemetry, None)


def test_negotiation_route_returns_503_without_persisting_a_hidden_fallback(monkeypatch):
    from database import init_db
    import main
    from negotiation_coach import NegotiationCoachUnavailable

    init_db()
    client = TestClient(main.app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)

    def unavailable(_context):
        raise NegotiationCoachUnavailable("Negotiation coach unavailable.")

    monkeypatch.setattr(main, "coach_negotiation", unavailable)
    response = client.post(
        f"/api/applications/workspaces/{workspace['id']}/negotiation/rehearse",
        headers=headers,
        json={
            "priorities": ["Role scope"],
            "scenario": "The recruiter asks for my minimum.",
        },
    )

    assert response.status_code == 503
    assert response.json() == {"detail": "Negotiation coach unavailable."}
    reloaded = client.get(
        f"/api/applications/workspaces/{workspace['id']}",
        headers=headers,
    )
    assert "negotiation" not in reloaded.json()["role_metadata"]


def test_negotiation_route_rejects_empty_or_duplicate_priorities():
    from database import init_db
    from main import app

    init_db()
    client = TestClient(app)
    headers = _signup(client)
    workspace = _create_workspace_with_resume(client, headers)
    path = f"/api/applications/workspaces/{workspace['id']}/negotiation/rehearse"

    for priorities in (["  "], ["Role scope", "role scope"]):
        response = client.post(
            path,
            headers=headers,
            json={
                "priorities": priorities,
                "scenario": "The recruiter asks for my minimum.",
            },
        )
        assert response.status_code == 422


def test_negotiation_coach_fails_closed_on_unavailable_or_invented_figures(monkeypatch):
    import negotiation_coach

    context = {
        "company": "Example Semiconductor",
        "role": "Manufacturing Manager",
        "scenario": "The employer offered S$9,000 monthly base.",
        "priorities": ["Role scope"],
        "anchor_options": [{"value": "S$9,000", "definition": "monthly base"}],
    }
    valid = {
        "opening": [
            "Thank them for the offer.",
            "Confirm that S$9,000 means monthly base before responding.",
        ],
        "questions": ["Which package definition applies?"],
        "priority_order": ["Role scope"],
    }
    monkeypatch.setattr(negotiation_coach, "call_sealion_json", lambda **_kwargs: __import__("json").dumps(valid))
    coaching = negotiation_coach.coach_negotiation(context)
    assert coaching["opening"] == "Thank them for the offer. Confirm that S$9,000 means monthly base before responding."
    assert coaching["trade_offs"] == ["Protect Role scope before trading another term."]

    invented = {**valid, "opening": "Counter at S$12,000 monthly base."}
    monkeypatch.setattr(
        negotiation_coach,
        "call_sealion_json",
        lambda **_kwargs: __import__("json").dumps(invented),
    )
    try:
        negotiation_coach.coach_negotiation(context)
    except negotiation_coach.NegotiationCoachUnavailable as error:
        assert "unsupported figures" in str(error)
    else:
        raise AssertionError("invented compensation figure was accepted")

    invented_period = {**valid, "opening": "Ask for a review after six months."}
    monkeypatch.setattr(
        negotiation_coach,
        "call_sealion_json",
        lambda **_kwargs: __import__("json").dumps(invented_period),
    )
    try:
        negotiation_coach.coach_negotiation(context)
    except negotiation_coach.NegotiationCoachUnavailable as error:
        assert "unsupported figures" in str(error)
    else:
        raise AssertionError("invented review period was accepted")

    monkeypatch.setattr(negotiation_coach, "call_sealion_json", lambda **_kwargs: None)
    try:
        negotiation_coach.coach_negotiation(context)
    except negotiation_coach.NegotiationCoachUnavailable as error:
        assert "unavailable" in str(error)
    else:
        raise AssertionError("unavailable coach silently fell back")
