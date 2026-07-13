#!/usr/bin/env python3
"""
Feature test suite for Job Hunter SG backend.
Run: cd backend && python -m pytest tests/test_features.py -v
"""
from __future__ import annotations

import os
import sys
from itertools import chain, repeat
from types import SimpleNamespace

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


# ═══════════════════════════════════════════════════════════════════════════
# 1. Skill Extractor
# ═══════════════════════════════════════════════════════════════════════════


class TestSkillExtractor:
    """Tests for skill_extractor.py"""

    def test_extract_known_skills_from_jd(self):
        """Should find multi-word skills that exist in KNOWN_SKILLS."""
        from skill_extractor import extract_skill_phrases

        jd = (
            "We need someone with experience in machine learning, "
            "project management, and data analysis. "
            "Cloud computing skills preferred."
        )
        skills = extract_skill_phrases(jd)
        skill_set = set(skills)
        assert "machine learning" in skill_set
        assert "project management" in skill_set
        assert "data analysis" in skill_set
        assert "cloud computing" in skill_set

    def test_no_skills_from_empty_text(self):
        from skill_extractor import extract_skill_phrases

        assert extract_skill_phrases("") == []
        assert extract_skill_phrases("   ") == []

    def test_metadata_skills_in_jd_text_are_included(self):
        """If a metadata skill appears in the JD text, always include it."""
        from skill_extractor import extract_skill_phrases

        jd = "This role requires strong quality management experience."
        skills = extract_skill_phrases(jd, job_skills=["Quality Management"])
        assert "quality management" in skills

    def test_metadata_skills_not_in_jd_and_not_known_are_excluded(self):
        """Metadata skills that don't appear in JD and aren't known should be filtered."""
        from skill_extractor import extract_skill_phrases

        jd = "We need a program manager with data analysis skills."
        skills = extract_skill_phrases(jd, job_skills=["Medical Study", "Basket Weaving"])
        assert "medical study" not in skills
        assert "basket weaving" not in skills

    def test_metadata_skills_known_but_not_in_jd_are_included(self):
        """Metadata skills in KNOWN_SKILLS but not in JD text should still be included."""
        from skill_extractor import extract_skill_phrases

        jd = "We need a software engineer."
        skills = extract_skill_phrases(jd, job_skills=["Machine Learning"])
        assert "machine learning" in skills

    def test_ngram_extraction(self):
        """The n-gram tokenizer should produce clean bigrams/trigrams."""
        from skill_extractor import _tokenize_for_ngrams, _extract_ngrams

        tokens = _tokenize_for_ngrams("strong project management skills required")
        ngrams = _extract_ngrams(tokens, 2)
        assert "project management" in ngrams
        # "skills required" should be filtered (stop-word-heavy)
        assert "skills required" not in ngrams

    def test_stop_words_filtered_from_ngrams(self):
        """N-grams starting or ending with stop words should be excluded."""
        from skill_extractor import _tokenize_for_ngrams, _extract_ngrams

        tokens = _tokenize_for_ngrams("the ability to work in a team")
        bigrams = _extract_ngrams(tokens, 2)
        # "the ability" starts with stop word
        assert not any(bg.startswith("the ") for bg in bigrams)
        # "a team" starts with stop word
        assert not any(bg.startswith("a ") for bg in bigrams)

    def test_resume_skill_matching(self):
        """Should correctly match/miss skills against resume text."""
        from skill_extractor import match_resume_skills

        resume = "Led machine learning projects and managed cross-functional teams."
        jd_skills = ["machine learning", "project management", "data analysis"]
        result = match_resume_skills(resume, jd_skills)

        matched_skills = [m["skill"] for m in result["matched"]]
        missing_skills = [m["skill"] for m in result["missing"]]
        assert "machine learning" in matched_skills
        assert "data analysis" in missing_skills
        assert result["match_percent"] > 0

    def test_resume_matching_empty_skills(self):
        from skill_extractor import match_resume_skills

        result = match_resume_skills("some resume text", [])
        assert result["matched"] == []
        assert result["missing"] == []

    def test_canonical_ats_terms_include_single_word_tech_and_filter_noise(self):
        from ats_terms import build_job_ats_terms

        jd = (
            "Build data pipelines in Python and SQL for semiconductor manufacturing. "
            "Medical Study is listed in the source taxonomy but should not be surfaced."
        )
        parsed_jd = {
            "required_skills": ["data pipelines", "semiconductor manufacturing"],
            "preferred_skills": [],
            "single_word_skills": ["python", "sql"],
            "competency_signals": {},
        }
        terms = build_job_ats_terms(
            jd_text=jd,
            job_skills=["Medical Study", "Python"],
            parsed_jd=parsed_jd,
            job_title="Data Engineer",
        )
        skills = {item["skill"] for item in terms}
        assert "python" in skills
        assert "sql" in skills
        assert "semiconductor manufacturing" in skills
        assert "medical study" not in skills

    def test_resume_scorer_keyword_match_uses_canonical_ats_terms(self):
        from resume_scorer import ResumeScorer

        scorer = ResumeScorer()
        result = scorer.analyze(
            "Built Python automation and SQL reporting for wafer fabrication teams.",
            "Role requires Python, SQL, and wafer fabrication experience.",
        )
        matched = {item["skill"] for item in result["keyword_match"]["matched"]}
        assert "python" in matched
        assert "sql" in matched
        assert "wafer fabrication" in matched

    def test_canonical_job_terms_include_parsed_jd_and_single_word_terms(self):
        from main import _build_canonical_job_terms

        job = SimpleNamespace(
            title="Principal Program Manager, FE Strategy & Operations",
            description=(
                "Drive wafer fabrication analytics, supply chain management, and "
                "lead implementation of Industry 4.0 initiatives with Python and SQL."
            ),
            skills=[
                "Wafer Fabrication",
                "Supply Chain Management",
                "Communication Skills",
            ],
            parsed_jd={
                "required_skills": ["wafer fabrication", "supply chain management"],
                "preferred_skills": ["lead implementation"],
                "single_word_skills": ["python", "sql"],
            },
        )

        terms = _build_canonical_job_terms(job, None)
        labels = {item["skill"].lower() for item in terms}

        assert "wafer fabrication" in labels
        assert "supply chain management" in labels
        assert "lead implementation" in labels
        assert "python" in labels
        assert "sql" in labels

    def test_canonical_job_terms_filter_noise_and_dedupe(self):
        from main import _build_canonical_job_terms

        job = SimpleNamespace(
            title="Senior Engineer",
            description="Build machine learning pipelines and quality systems.",
            skills=["Professional Experience", "Machine Learning", "machine learning"],
            parsed_jd={
                "required_skills": ["machine learning"],
                "preferred_skills": ["quality systems"],
                "single_word_skills": ["AI"],
            },
        )

        terms = _build_canonical_job_terms(job, None)
        labels = [item["skill"].lower() for item in terms]

        assert "professional experience" not in labels
        assert labels.count("machine learning") == 1
        assert "quality systems" in labels
        assert "ai" in labels

    def test_careersgov_jd_surfaces_technical_cues_without_source_tags(self):
        from ats_terms import build_job_ats_terms
        from jd_preparser import preparse_job_description

        jd = (
            "We are currently seeking a Modelling and Simulation Lead with a focus on "
            "End-to-End Systems to analyse the performance of enterprise applications "
            "across various communication networks. Your expertise in end-to-end system "
            "modelling and simulation techniques will ensure a thorough understanding "
            "of application performance metrics, such as response time, latency, and "
            "error rates. Proficiency in end-to-end system modelling and simulation "
            "tools, such as OPNET, NS-3, QuNetSim, or similar."
        )
        parsed = preparse_job_description(
            jd,
            skills=[],
            job_title="Lead Engineer (Modeling and Simulation)",
        )
        terms = build_job_ats_terms(
            jd_text=jd,
            job_skills=[],
            parsed_jd=parsed,
            job_title="Lead Engineer (Modeling and Simulation)",
        )
        labels = {item["skill"].lower() for item in terms}

        assert "modeling and simulation" in labels or "modelling and simulation" in labels
        assert "communication networks" in labels
        assert "end-to-end systems" in labels or "end-to-end system modelling" in labels
        assert any(tool in labels for tool in {"opnet", "ns-3", "qunetsim"})

    def test_careersgov_html_cleanup_preserves_bullets_and_paragraphs(self):
        from scraper import _clean_html

        html = "<p>Overview paragraph.</p><ul><li>First item</li><li>Second item</li></ul>"
        cleaned = _clean_html(html)

        assert "Overview paragraph." in cleaned
        assert "\n- First item" in cleaned
        assert "\n- Second item" in cleaned

    def test_careersgov_detail_uses_cached_json_dump(self, monkeypatch):
        from scraper import CareersGovScraper

        monkeypatch.setattr(
            CareersGovScraper,
            "_fetch_data",
            classmethod(
                lambda cls: [
                    {
                        "jobId": "123",
                        "postingNo": "ABC",
                        "jobTitle": "Data Engineer",
                        "agency": "GovTech",
                        "jobDescription": "<p>Build pipelines.</p>",
                        "jobResponsibilities": "<ul><li>Own data quality</li></ul>",
                        "jobRequirements": "Python and SQL",
                    }
                ]
            ),
        )

        detail = CareersGovScraper().get_job_detail("/jobs/hrp/123/ABC")

        assert detail["companyName"] == "GovTech"
        assert "Build pipelines." in detail["jobDescription"]
        assert "- Own data quality" in detail["jobDescription"]
        assert "Python and SQL" in detail["jobDescription"]

    def test_ats_terms_exclude_benefits_bullets_but_keep_real_skills(self):
        from ats_terms import build_job_ats_terms

        jd = (
            "Benefits:\n"
            "• Generous PTO\n"
            "• Health Benefits\n"
            "• Emotional Health Resources\n"
            "• HeyGen Offers\n\n"
            "Requirements:\n"
            "• Machine Learning\n"
            "• Computer Vision\n"
            "• Python\n"
        )

        terms = build_job_ats_terms(
            jd_text=jd,
            job_skills=[],
            parsed_jd={"required_skills": [], "preferred_skills": [], "single_word_skills": ["python"]},
            job_title="ML Engineer",
        )
        labels = {item["skill"].lower() for item in terms}

        assert "machine learning" in labels
        assert "computer vision" in labels
        assert "python" in labels
        assert "generous pto" not in labels
        assert "health benefits" not in labels
        assert "emotional health resources" not in labels
        assert "heygen offers" not in labels

    def test_ats_terms_drop_or_fragment_but_keep_individual_skills(self):
        from ats_terms import build_job_ats_terms

        terms = build_job_ats_terms(
            jd_text="Requirements: Python or C++",
            job_skills=[],
            parsed_jd={"required_skills": [], "preferred_skills": [], "single_word_skills": ["python", "c++"]},
            job_title="Software Engineer",
        )
        labels = {item["skill"].lower() for item in terms}

        assert "python" in labels
        assert "c++" in labels
        assert "python or c++" not in labels


# ═══════════════════════════════════════════════════════════════════════════
# 2. Resume Parser
# ═══════════════════════════════════════════════════════════════════════════


class TestResumeParser:
    """Tests for resume_parser.py"""

    def test_validate_upload_pdf(self):
        from resume_parser import validate_upload

        result = validate_upload("resume.pdf", "application/pdf", 1024)
        assert result == "pdf"

    def test_validate_upload_docx(self):
        from resume_parser import validate_upload

        result = validate_upload(
            "resume.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            1024,
        )
        assert result == "docx"

    def test_validate_upload_rejects_too_large(self):
        from resume_parser import validate_upload

        with pytest.raises(ValueError, match="too large"):
            validate_upload("resume.pdf", "application/pdf", 10 * 1024 * 1024)

    def test_validate_upload_rejects_empty(self):
        from resume_parser import validate_upload

        with pytest.raises(ValueError, match="empty"):
            validate_upload("resume.pdf", "application/pdf", 0)

    def test_validate_upload_rejects_unsupported(self):
        from resume_parser import validate_upload

        with pytest.raises(ValueError, match="Unsupported"):
            validate_upload("resume.txt", "text/plain", 1024)

    def test_join_broken_lines_hyphenated(self):
        """Hyphenated line breaks should be joined."""
        from resume_parser import _join_broken_lines

        text = "Led cross-\nfunctional delivery"
        result = _join_broken_lines(text)
        assert result == "Led cross-functional delivery"

    def test_join_broken_lines_preserves_pdf_bullets_without_spaces(self):
        from resume_parser import _join_broken_lines, _parse_quality

        text = (
            "Selected for an industry project targeting a 90% improvement.\n"
            "•Built the production agent scaffold.\n"
            "•Designed the validation workflow."
        )

        result = _join_broken_lines(text)

        assert result.splitlines() == [
            "Selected for an industry project targeting a 90% improvement.",
            "• Built the production agent scaffold.",
            "• Designed the validation workflow.",
        ]
        assert _parse_quality(result, "pdf")["signals"]["bullet_line_count"] == 2

    def test_join_broken_lines_preserves_meaningful_hyphen(self):
        from resume_parser import _join_broken_lines

        assert _join_broken_lines("Produced evidence-\nbacked reports") == (
            "Produced evidence-backed reports"
        )

    def test_join_broken_lines_preserves_sections(self):
        """Section headers should stay on their own lines."""
        from resume_parser import _join_broken_lines

        text = "PROFESSIONAL EXPERIENCE\nSoftware Engineer at Google"
        result = _join_broken_lines(text)
        lines = [l for l in result.split("\n") if l.strip()]
        assert lines[0].strip() == "PROFESSIONAL EXPERIENCE"

    def test_join_broken_lines_preserves_executive_summary_heading(self):
        """Executive Summary should not be merged into the paragraph below."""
        from resume_parser import _join_broken_lines

        text = (
            "EXECUTIVE SUMMARY\n"
            "Transformation leader with 7+ years in semiconductor manufacturing.\n"
            "Led cross-functional programs across global fabs."
        )
        result = _join_broken_lines(text)
        lines = [l for l in result.split("\n") if l.strip()]
        assert lines[0].strip() == "EXECUTIVE SUMMARY"
        assert lines[1].startswith("Transformation leader")

    def test_join_broken_lines_does_not_guess_semantic_paragraphs(self):
        from resume_parser import _join_broken_lines

        text = (
            "CORE SKILLS\n"
            "Leadership and Delivery: programme management and adoption\n"
            "Agentic AI and LLM Engineering: LangGraph and RAG"
        )

        assert _join_broken_lines(text) == text

    def test_name_detection(self):
        """parse_resume metadata should extract a name from the first lines."""
        from resume_parser import parse_resume
        import io

        # Create a minimal DOCX for testing
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("JOHN SMITH")
            doc.add_paragraph("john@example.com")
            doc.add_paragraph("EXPERIENCE")
            doc.add_paragraph("Software Engineer at Google")
            buf = io.BytesIO()
            doc.save(buf)
            result = parse_resume("test.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", buf.getvalue())
            assert result["name"] == "JOHN SMITH"
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_docx_list_bullets_preserved(self):
        """DOCX list bullets should be extracted as visible bullet lines."""
        from resume_parser import parse_resume
        import io

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Jane Doe")
            doc.add_paragraph("PROFESSIONAL EXPERIENCE")
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run("Led cross-functional process integration across 3 fabs")
            buf = io.BytesIO()
            doc.save(buf)

            result = parse_resume(
                "bullets.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                buf.getvalue(),
            )
            assert "• Led cross-functional process integration across 3 fabs" in result["text"]
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_quality_warns_on_low_signal_docx(self):
        """Upload metadata should flag extracts that look too short to trust."""
        from resume_parser import parse_resume
        import io

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Jane Doe")
            doc.add_paragraph("jane@example.com")
            doc.add_paragraph("Experienced operator")
            buf = io.BytesIO()
            doc.save(buf)

            result = parse_resume(
                "thin.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                buf.getvalue(),
            )
            quality = result["parse_quality"]
            assert quality["label"] == "review"
            assert quality["warnings"]
            assert quality["signals"]["word_count"] < 120
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_quality_good_for_structured_docx(self):
        """A normal structured resume should not show parse warnings."""
        from resume_parser import parse_resume
        import io

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Jane Doe")
            doc.add_paragraph("jane@example.com | +65 9123 4567")
            doc.add_paragraph("SUMMARY")
            doc.add_paragraph("Engineering leader with experience across manufacturing and analytics.")
            doc.add_paragraph("PROFESSIONAL EXPERIENCE")
            doc.add_paragraph("Manager | Micron Technology | 2022-2025")
            for _ in range(12):
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run("Led cross-functional process improvements across 3 fabs, improving cycle time by 20%")
            doc.add_paragraph("EDUCATION")
            doc.add_paragraph("Bachelor of Engineering, National University of Singapore")
            doc.add_paragraph("SKILLS")
            doc.add_paragraph("Python, SQL, Project Management, Data Analysis")
            buf = io.BytesIO()
            doc.save(buf)

            result = parse_resume(
                "structured.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                buf.getvalue(),
            )
            assert result["parse_quality"]["label"] == "good"
            assert result["parse_quality"]["warnings"] == []
        except ImportError:
            pytest.skip("python-docx not installed")


# ═══════════════════════════════════════════════════════════════════════════
# 3. Resume Scorer
# ═══════════════════════════════════════════════════════════════════════════


class TestResumeScorer:
    """Tests for resume_scorer.py"""

    def test_scorer_returns_all_dimensions(self):
        from resume_scorer import ResumeScorer

        scorer = ResumeScorer()
        result = scorer.analyze(
            resume_text="Led machine learning projects at Google. Managed a team of 10 engineers. "
            "Delivered $5M in cost savings through automation. Built CI/CD pipelines.",
            job_description="",
        )
        assert "overall_score" in result
        assert "dimensions" in result
        assert result["overall_score"] >= 0

    def test_scorer_with_jd_returns_keyword_match(self):
        from resume_scorer import ResumeScorer

        scorer = ResumeScorer()
        result = scorer.analyze(
            resume_text="Experienced in machine learning, data analysis, and project management. "
            "Led cross-functional teams. Built data pipelines.",
            job_description="Looking for machine learning engineer with data analysis and cloud computing skills.",
        )
        assert "keyword_match" in result
        kw = result["keyword_match"]
        assert "matched" in kw
        assert "missing" in kw

    def test_competencies_include_matched_and_missing_keywords(self):
        """Each competency item should have matched_keywords and missing_keywords."""
        from resume_scorer import ResumeScorer

        scorer = ResumeScorer()
        result = scorer.analyze(
            resume_text="Collaborated with teams. Led initiatives. Analyzed data. "
            "Managed stakeholders. Presented findings.",
        )
        competencies = result.get("dimensions", {}).get("competencies", {})
        if competencies and "items" in competencies:
            for name, item in competencies["items"].items():
                assert "matched_keywords" in item, f"{name} missing matched_keywords"
                assert "missing_keywords" in item, f"{name} missing missing_keywords"
                assert isinstance(item["matched_keywords"], list)
                assert isinstance(item["missing_keywords"], list)

    def test_action_verb_scoring(self):
        """Bullets starting with action verbs should score higher."""
        from resume_scorer import ResumeScorer

        scorer = ResumeScorer()
        strong = scorer.analyze(
            resume_text=(
                "EXPERIENCE\n"
                "• Led a team of 10 engineers to deliver $5M platform\n"
                "• Designed microservices architecture for 1M users\n"
                "• Implemented CI/CD pipeline reducing deploy time by 40%\n"
                "• Managed cross-functional collaboration across 3 offices\n"
            ),
        )
        weak = scorer.analyze(
            resume_text=(
                "EXPERIENCE\n"
                "• Was responsible for team management\n"
                "• Helped with architecture design\n"
                "• Assisted in pipeline setup\n"
                "• Various duties as assigned\n"
            ),
        )
        assert strong["overall_score"] > weak["overall_score"]

    def test_empty_resume_returns_low_score(self):
        from resume_scorer import ResumeScorer

        scorer = ResumeScorer()
        result = scorer.analyze(resume_text="")
        assert result["overall_score"] < 50


# ═══════════════════════════════════════════════════════════════════════════
# 4. Sanitizer
# ═══════════════════════════════════════════════════════════════════════════


class TestSanitizer:
    """Tests for sanitizer.py"""

    def test_sanitize_strips_html(self):
        from sanitizer import sanitize_user_input

        result = sanitize_user_input("<script>alert('xss')</script>Hello")
        assert "<script>" not in result
        assert "Hello" in result

    def test_sanitize_resume_text(self):
        from sanitizer import sanitize_resume_text

        result = sanitize_resume_text("  Normal resume text  \n  With lines  ")
        assert result.strip() == result  # Should be stripped
        assert result == "Normal resume text\nWith lines"

    def test_sanitize_resume_text_preserves_resume_sections(self):
        from sanitizer import sanitize_resume_text

        raw = (
            "Haoming Koo\n"
            "haomingkoo@gmail.com\n\n"
            "PROFESSIONAL SUMMARY\n"
            "Transformation leader.\n\n"
            "EXPERIENCE\n"
            "Led delivery.\n"
        )
        result = sanitize_resume_text(raw)
        assert "PROFESSIONAL SUMMARY\nTransformation leader." in result
        assert "\n\nEXPERIENCE\n" in result

    def test_sanitize_job(self):
        from sanitizer import sanitize_job

        job = {
            "title": "<b>Engineer</b>",
            "company": "Google<script>",
            "description": "Normal description",
            "dedup_key": "abc123",
        }
        result = sanitize_job(job)
        assert "<b>" not in result["title"]
        assert "<script>" not in result["company"]


# ═══════════════════════════════════════════════════════════════════════════
# 5. Auth
# ═══════════════════════════════════════════════════════════════════════════


class TestAuth:
    """Tests for auth.py"""

    def test_password_hashing(self):
        from auth import hash_password, verify_password

        pw = "SecurePass123!"
        hashed = hash_password(pw)
        assert hashed != pw
        assert verify_password(pw, hashed)
        assert not verify_password("wrong", hashed)

    def test_password_validation(self):
        from auth import validate_password
        from fastapi import HTTPException

        # Too short — raises HTTPException, not ValueError
        with pytest.raises(HTTPException):
            validate_password("short")

        # Valid — should not raise
        validate_password("LongEnoughPassword123")

    def test_token_creation(self):
        from auth import create_token

        token = create_token(user_id=1)
        assert isinstance(token, str)
        assert len(token) > 0


# ═══════════════════════════════════════════════════════════════════════════
# 6. API Endpoints (integration)
# ═══════════════════════════════════════════════════════════════════════════


class TestAPIEndpoints:
    """Integration tests for FastAPI endpoints."""

    @pytest.fixture
    def client(self):
        from fastapi.testclient import TestClient
        from main import app
        return TestClient(app)

    def test_health_endpoint(self, client):
        resp = client.get("/api/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ok"

    def test_jobs_endpoint(self, client):
        resp = client.get("/api/jobs")
        assert resp.status_code == 200
        data = resp.json()
        # Returns paginated response with jobs list
        assert "jobs" in data or isinstance(data, list)

    def test_parse_job_posted_at_prefers_newer_relative_dates(self):
        from main import _parse_job_posted_at

        recent = _parse_job_posted_at("Posted 4 Days Ago")
        stale = _parse_job_posted_at("Posted 30+ Days Ago")
        assert recent > stale

    def test_trending_skills_endpoint(self, client):
        resp = client.get("/api/skills/trending")
        assert resp.status_code == 200
        assert isinstance(resp.json(), list)

    def test_analytics_trends_endpoint(self, client):
        from main import _ANALYTICS_CACHE_LOCK, _analytics_query_cache, _clear_analytics_cache

        _clear_analytics_cache()
        resp = client.get("/api/analytics/trends?weeks=4")
        assert resp.status_code == 200
        data = resp.json()
        assert "series" in data
        assert "recent_top_titles" in data
        assert "recent_ats_terms" in data
        second = client.get("/api/analytics/trends?weeks=4")
        assert second.json() == data
        with _ANALYTICS_CACHE_LOCK:
            assert any(key[0] == "trends" and key[-1] == 4 for key in _analytics_query_cache)

    def test_resume_score_requires_text(self, client):
        resp = client.post("/api/resume/score", json={
            "resume_text": "",
            "job_description": "",
        })
        # Should return 200 with zero score or 422 validation error
        assert resp.status_code in (200, 422)

    def test_admin_seed_rejects_bad_key(self, client):
        resp = client.post(
            "/api/admin/seed",
            json={"sources": "mcf", "limit": 1},
            headers={"Authorization": "Bearer wrong_key"},
        )
        assert resp.status_code == 403

    def test_signup_and_login_flow(self, client, monkeypatch):
        import secrets
        import main
        email = f"test_{secrets.token_hex(4)}@aisg.sg"
        pw = "TestPassword123!"
        sent = {}
        monkeypatch.setattr(main, "email_configured", lambda: True)
        monkeypatch.setattr(
            main,
            "_send_verification_email",
            lambda _user, token: sent.setdefault("token", token),
        )

        # Signup
        resp = client.post("/api/auth/signup", json={
            "email": email,
            "password": pw,
            "name": "Test User",
            "accepted_terms": True,
        })
        assert resp.status_code == 200
        assert "token" not in resp.json()

        resp = client.post(
            "/api/auth/verify-email",
            json={
                "token": sent["token"],
                "password": pw,
                "name": "Test User",
                "accepted_terms": True,
            },
        )
        assert resp.status_code == 200
        assert resp.json()["user"]["email"] == email
        assert resp.json()["user"]["tier"] == "user"

        # Login
        resp = client.post("/api/auth/login", json={
            "email": email,
            "password": pw,
        })
        assert resp.status_code == 200
        assert "token" in resp.json()

    def test_tracked_status_history_is_append_only(self, client, monkeypatch):
        import secrets
        import main
        from database import init_db

        init_db()

        email = f"tracked_{secrets.token_hex(4)}@aisg.sg"
        pw = "TestPassword123!"
        sent = {}
        monkeypatch.setattr(main, "email_configured", lambda: True)
        monkeypatch.setattr(
            main,
            "_send_verification_email",
            lambda _user, token: sent.setdefault("token", token),
        )
        signup = client.post("/api/auth/signup", json={
            "email": email,
            "password": pw,
            "name": "Tracked User",
            "accepted_terms": True,
        })
        assert signup.status_code == 200
        verification = client.post(
            "/api/auth/verify-email",
            json={
                "token": sent["token"],
                "password": pw,
                "name": "Test User",
                "accepted_terms": True,
            },
        )
        assert verification.status_code == 200
        headers = {"Authorization": f"Bearer {verification.json()['token']}"}

        created = client.post("/api/tracked", json={
            "company": "Example Co",
            "role": "AI Program Manager",
            "date_applied": "2026-07-03",
            "status": "saved",
        }, headers=headers)
        assert created.status_code == 201
        tracked = created.json()
        assert [event["stage"] for event in tracked["stage_history"]] == ["saved"]

        updated = client.put(
            f"/api/tracked/{tracked['id']}",
            json={"status": "screening"},
            headers=headers,
        )
        assert updated.status_code == 200
        assert [event["stage"] for event in updated.json()["stage_history"]] == [
            "saved",
            "screening",
        ]

    def test_static_frontend_or_no_static(self, client):
        """/ should serve frontend if static dir exists, or 404 otherwise."""
        resp = client.get("/")
        # 200 if static dir exists, 404 if not (local dev without build)
        assert resp.status_code in (200, 307, 404)


class TestBackfillProgress:
    def test_preview_backfill_reports_live_rate_and_eta(self, monkeypatch):
        from backfill_enrichment import backfill_previews

        jobs = [
            SimpleNamespace(
                id=1,
                title="Data Scientist",
                description="Build Python pipelines.",
                skills=["Python"],
                parsed_jd=None,
                job_terms_preview=None,
                salary="",
                company="AISG",
                agency="",
            ),
            SimpleNamespace(
                id=2,
                title="ML Engineer",
                description="Deploy ML systems.",
                skills=["Machine Learning"],
                parsed_jd=None,
                job_terms_preview=None,
                salary="",
                company="AISG",
                agency="",
            ),
            SimpleNamespace(
                id=3,
                title="Analytics Engineer",
                description="Model analytics data.",
                skills=["SQL"],
                parsed_jd=None,
                job_terms_preview=None,
                salary="",
                company="AISG",
                agency="",
            ),
        ]

        class FakeQuery:
            def __init__(self, items, start=0, size=None):
                self.items = items
                self.start = start
                self.size = size

            def filter(self, *_args, **_kwargs):
                return self

            def order_by(self, *_args, **_kwargs):
                return self

            def count(self):
                return len(self.items)

            def offset(self, amount):
                return FakeQuery(self.items, start=amount, size=self.size)

            def limit(self, amount):
                return FakeQuery(self.items, start=self.start, size=amount)

            def all(self):
                items = self.items[self.start:]
                if self.size is not None:
                    items = items[:self.size]
                return items

        class FakeDB:
            def __init__(self, items):
                self.items = items
                self.closed = False
                self.commits = 0

            def query(self, _model):
                return FakeQuery(self.items)

            def commit(self):
                self.commits += 1

            def close(self):
                self.closed = True

        fake_db = FakeDB(jobs)
        progress_updates = []
        timeline = iter(chain([0, 60, 120, 180, 180], repeat(180)))

        monkeypatch.setattr("backfill_enrichment.SessionLocal", lambda: fake_db)
        monkeypatch.setattr("backfill_enrichment.preparse_jd", lambda *args, **kwargs: {"required_skills": []})
        monkeypatch.setattr("backfill_enrichment.build_job_ats_terms", lambda **kwargs: [{"skill": "python"}])
        monkeypatch.setattr("backfill_enrichment.analyze_job_description", lambda **kwargs: {"fit": "good"})
        monkeypatch.setattr("backfill_enrichment.time.time", lambda: next(timeline))

        processed = backfill_previews(
            batch_size=1,
            progress_callback=lambda **kwargs: progress_updates.append(kwargs),
            refresh_preview=True,
        )

        assert processed == 3
        assert fake_db.closed is True
        assert fake_db.commits == 3
        assert progress_updates[0]["preview_total"] == 3
        assert progress_updates[0]["rate_per_min"] == 0.0
        assert progress_updates[0]["eta_minutes"] == 0.0

        live_updates = [
            update
            for update in progress_updates
            if 0 < update.get("preview_done", 0) < update.get("preview_total", 0)
        ]
        assert live_updates
        assert all(update["rate_per_min"] > 0 for update in live_updates)
        assert all(update["eta_minutes"] > 0 for update in live_updates)

        assert progress_updates[-1]["preview_done"] == 3
        assert progress_updates[-1]["preview_total"] == 3
        assert progress_updates[-1]["eta_minutes"] == 0.0
        assert all(job.job_terms_preview == ["Python"] for job in jobs)

    def test_backfill_cli_passes_reparse_flag_to_preview_phase(self, monkeypatch):
        import argparse
        import backfill_enrichment

        captured = {}

        monkeypatch.setattr(
            argparse.ArgumentParser,
            "parse_args",
            lambda self: SimpleNamespace(
                preview_only=True,
                summary_only=False,
                summary_limit=0,
                batch_size=50,
                refresh_preview=True,
                reparse=True,
            ),
        )
        monkeypatch.setattr(backfill_enrichment, "init_db", lambda: None)
        def fake_backfill_previews(**kwargs):
            captured["preview"] = kwargs
            return 0
        monkeypatch.setattr(
            backfill_enrichment,
            "backfill_previews",
            fake_backfill_previews,
        )
        monkeypatch.setattr(
            backfill_enrichment,
            "backfill_summaries",
            lambda **kwargs: pytest.fail("summary phase should not run when preview_only=True"),
        )

        backfill_enrichment.main()

        assert captured["preview"]["batch_size"] == 50
        assert captured["preview"]["refresh_preview"] is True
        assert captured["preview"]["reparse"] is True

    def test_refresh_preview_uses_stable_ordered_pagination(self, monkeypatch):
        from backfill_enrichment import backfill_previews

        jobs = [
            SimpleNamespace(id=1, title="One", description="Desc 1", skills=[], parsed_jd=None, job_terms_preview=None, salary="", company="", agency=""),
            SimpleNamespace(id=2, title="Two", description="Desc 2", skills=[], parsed_jd=None, job_terms_preview=None, salary="", company="", agency=""),
            SimpleNamespace(id=3, title="Three", description="Desc 3", skills=[], parsed_jd=None, job_terms_preview=None, salary="", company="", agency=""),
            SimpleNamespace(id=4, title="Four", description="Desc 4", skills=[], parsed_jd=None, job_terms_preview=None, salary="", company="", agency=""),
        ]

        class FakeQuery:
            def __init__(self, items, start=0, size=None, ordered=False, order_calls=None, all_calls=None):
                self.items = items
                self.start = start
                self.size = size
                self.ordered = ordered
                self.order_calls = order_calls if order_calls is not None else {"count": 0}
                self.all_calls = all_calls if all_calls is not None else {"count": 0}

            def filter(self, *_args, **_kwargs):
                return self

            def count(self):
                return len(self.items)

            def order_by(self, *_args, **_kwargs):
                self.order_calls["count"] += 1
                return FakeQuery(
                    self.items,
                    start=self.start,
                    size=self.size,
                    ordered=True,
                    order_calls=self.order_calls,
                    all_calls=self.all_calls,
                )

            def offset(self, amount):
                return FakeQuery(
                    self.items,
                    start=amount,
                    size=self.size,
                    ordered=self.ordered,
                    order_calls=self.order_calls,
                    all_calls=self.all_calls,
                )

            def limit(self, amount):
                return FakeQuery(
                    self.items,
                    start=self.start,
                    size=amount,
                    ordered=self.ordered,
                    order_calls=self.order_calls,
                    all_calls=self.all_calls,
                )

            def all(self):
                self.all_calls["count"] += 1
                if self.ordered:
                    items = sorted(self.items, key=lambda item: item.id)
                else:
                    # Simulate unstable DB row order when no ORDER BY is present.
                    items = list(reversed(self.items)) if self.all_calls["count"] % 2 == 0 else list(self.items)
                items = items[self.start:]
                if self.size is not None:
                    items = items[:self.size]
                return items

        class FakeDB:
            def __init__(self, items):
                self.items = items

            def query(self, _model):
                return FakeQuery(self.items)

            def commit(self):
                pass

            def close(self):
                pass

        monkeypatch.setattr("backfill_enrichment.SessionLocal", lambda: FakeDB(jobs))
        monkeypatch.setattr("backfill_enrichment.preparse_jd", lambda *args, **kwargs: {"required_skills": []})
        monkeypatch.setattr("backfill_enrichment.build_job_ats_terms", lambda **kwargs: [{"skill": f"skill-{kwargs['job_title']}"}])
        monkeypatch.setattr("backfill_enrichment.analyze_job_description", lambda **kwargs: {"fit": "good"})

        processed = backfill_previews(batch_size=2, refresh_preview=True)

        assert processed == 4
        assert all(job.parsed_jd is not None for job in jobs)
        assert all(job.job_terms_preview is not None for job in jobs)
