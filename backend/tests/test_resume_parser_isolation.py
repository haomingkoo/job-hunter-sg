from __future__ import annotations

import io
from pathlib import Path
import sys
import time

import pytest
from fastapi.testclient import TestClient


PDF_TYPE = "application/pdf"
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


SAMPLE_LINES = [
    "Hui Shan Ang",
    "huishan@example.com | +65 9123 4567",
    "PROFESSIONAL SUMMARY",
    "Finance transformation leader with regional experience across accounting, analytics, automation, and business partnering.",
    "CORE SKILLS",
    "Financial reporting, Oracle Cloud, Power BI, Python, SQL, stakeholder management",
    "FINANCE PROCESS & TRANSFORMATION EXPERIENCE",
    "Finance Transformation Manager | Example Company | 2021 - Present",
    "- Led a regional close redesign across eight markets and reduced manual reconciliation time by 35%.",
    "- Implemented finance controls and reporting standards used by more than 100 business stakeholders.",
    "- Partnered with operations teams to simplify approval workflows and improve audit readiness.",
    "AUTOMATION & AI EXPERIENCE",
    "AI Finance Lead | Example Company | 2023 - Present",
    "- Built forecasting automation with Python and SQL, shortening monthly analysis by four working days.",
    "- Designed data quality checks that prevented incomplete records from reaching executive dashboards.",
    "- Trained finance users to review model evidence and escalate uncertain recommendations.",
    "EDUCATION & CERTIFICATIONS",
    "Bachelor of Accountancy, National University of Singapore, 2015",
    "LANGUAGES",
    "English, Mandarin",
]


def _sample_docx() -> bytes:
    from docx import Document

    document = Document()
    for line in SAMPLE_LINES:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _sample_pdf() -> bytes:
    from reportlab.pdfgen.canvas import Canvas

    buffer = io.BytesIO()
    canvas = Canvas(buffer)
    y = 800
    for line in SAMPLE_LINES:
        canvas.drawString(48, y, line)
        y -= 24
        if y < 48:
            canvas.showPage()
            y = 800
    canvas.save()
    return buffer.getvalue()


def _two_column_pdf() -> bytes:
    from reportlab.pdfgen.canvas import Canvas

    buffer = io.BytesIO()
    canvas = Canvas(buffer)
    canvas.drawString(48, 800, "PROFESSIONAL EXPERIENCE")
    y = 760
    for index in range(10):
        canvas.drawString(48, y, f"- Left result {index} improved monthly finance reporting")
        canvas.drawString(330, y, f"- Right result {index} supported regional finance users")
        y -= 32
    canvas.save()
    return buffer.getvalue()


def test_isolated_parser_returns_normal_pdf_result():
    from resume_document import is_resume_document
    from resume_parser import parse_resume_isolated

    path = Path(__file__).parents[1] / "templates/nus/NUS Guidelines.pdf"
    result = parse_resume_isolated(path.name, PDF_TYPE, path.read_bytes())

    assert result["file_type"] == "pdf"
    assert result["text"]
    assert is_resume_document(result["document"])
    assert result["document"]["raw_text"] == result["text"]
    assert result["document"]["source"]["format"] == "pdf"


def test_pdf_and_docx_adapters_produce_the_same_semantic_display_signature() -> None:
    from resume_parser import parse_resume

    pdf = parse_resume("resume.pdf", PDF_TYPE, _sample_pdf())["document"]
    docx = parse_resume("resume.docx", DOCX_TYPE, _sample_docx())["document"]

    def signature(document: dict) -> list[tuple[str, str, str]]:
        return [
            (block["kind"], block["section_key"], block["text"])
            for block in document["blocks"]
        ]

    assert signature(pdf) == signature(docx)
    assert [section["key"] for section in pdf["sections"]] == [
        section["key"] for section in docx["sections"]
    ]
    assert all(block["page"] == 1 for block in pdf["blocks"])
    assert any(block["style_name"] == "Normal" for block in docx["blocks"])


def test_isolated_parser_returns_normal_docx_result():
    from docx import Document
    from resume_parser import parse_resume_isolated

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com")
    document.add_paragraph("EXPERIENCE")
    document.add_paragraph("Software Engineer at Example")
    buffer = io.BytesIO()
    document.save(buffer)

    result = parse_resume_isolated("resume.docx", DOCX_TYPE, buffer.getvalue())

    assert result["file_type"] == "docx"
    assert result["name"] == "Jane Doe"
    assert result["document"]["raw_text"] == result["text"]


def test_docx_preserves_paragraph_and_table_order() -> None:
    from docx import Document
    from resume_parser import parse_resume

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("CORE SKILLS")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "SQL"
    merged = table.add_row().cells
    merged_cell = merged[0].merge(merged[1])
    merged_cell.text = "SELECTED PROJECTS"
    merged_cell.add_paragraph("Built an audit dashboard", style="List Bullet")
    document.add_paragraph("FINANCE TRANSFORMATION EXPERIENCE")
    document.add_paragraph("Finance Manager | Example Company | 2021 - Present")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_resume("resume.docx", DOCX_TYPE, buffer.getvalue())

    assert parsed["text"].index("CORE SKILLS") < parsed["text"].index("Python | SQL")
    assert parsed["text"].index("Python | SQL") < parsed["text"].index("FINANCE TRANSFORMATION EXPERIENCE")
    assert parsed["text"].count("SELECTED PROJECTS") == 1
    assert "SELECTED PROJECTS\n• Built an audit dashboard" in parsed["text"]


def test_docx_extracts_contact_from_word_header_without_treating_dates_as_phone() -> None:
    from docx import Document
    from resume_parser import parse_resume

    document = Document()
    document.sections[0].header.paragraphs[0].text = "Jane Doe | jane@example.com"
    document.add_paragraph("PROFESSIONAL EXPERIENCE")
    document.add_paragraph("Finance Manager | 2021 - 2024")
    document.add_paragraph("Led a regional reporting redesign used across eight markets.")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_resume("resume.docx", DOCX_TYPE, buffer.getvalue())

    assert parsed["text"].startswith("Jane Doe | jane@example.com")
    assert parsed["email"] == "jane@example.com"
    assert parsed["phone"] is None


@pytest.mark.parametrize(
    ("filename", "content_type", "file_bytes"),
    [
        ("resume.pdf", PDF_TYPE, _sample_pdf()),
        ("resume.docx", DOCX_TYPE, _sample_docx()),
    ],
)
def test_upload_preserves_custom_sections_end_to_end(
    filename: str,
    content_type: str,
    file_bytes: bytes,
) -> None:
    from main import app
    from resume_structurer import structure_resume

    response = TestClient(app).post(
        "/api/resume/upload",
        files={"file": (filename, file_bytes, content_type)},
    )

    assert response.status_code == 200
    parsed = response.json()
    assert parsed["parse_quality"]["signals"]["section_count"] == 6
    assert parsed["parse_quality"]["label"] == "good"
    assert parsed["email"] == "huishan@example.com"

    structured = structure_resume(parsed["text"])
    assert [section["key"] for section in structured["sections"]] == [
        "summary",
        "skills",
        "experience",
        "experience",
        "education",
        "languages",
    ]


def test_two_column_pdf_reports_layout_risk_instead_of_false_good() -> None:
    from resume_parser import parse_resume

    parsed = parse_resume("columns.pdf", PDF_TYPE, _two_column_pdf())

    assert parsed["parse_quality"]["label"] != "good"
    assert parsed["parse_quality"]["signals"]["possible_multi_column_layout"] is True
    assert any("columns" in warning for warning in parsed["parse_quality"]["warnings"])


def test_pasted_text_and_heading_confirmation_use_canonical_interface() -> None:
    from main import app
    from resume_document import is_resume_document

    client = TestClient(app)
    response = client.post(
        "/api/resume/ingest-text",
        json={
            "resume_text": "EXPERIENCE\n- Built a reporting platform.\n\nSELECTED TALKS\nSpoke at PyCon Singapore."
        },
    )
    assert response.status_code == 200
    document = response.json()
    assert is_resume_document(document)
    assert document["source"]["format"] == "text"
    candidate = document["heading_candidates"][0]

    confirmed = client.post(
        "/api/resume/confirm-heading",
        json={
            "document": document,
            "block_id": candidate["block_id"],
            "expected_revision": document["revision"],
            "section_key": None,
        },
    )
    assert confirmed.status_code == 200
    updated = confirmed.json()
    assert updated["raw_text"] == document["raw_text"]
    assert updated["sections"][-1]["label"] == "SELECTED TALKS"


@pytest.mark.parametrize(
    ("path", "content_type"),
    [
        (Path("/tmp/jobhunter-redteam-flate-bomb.pdf"), PDF_TYPE),
        (Path("/tmp/jobhunter-redteam-docx-dom-bomb.docx"), DOCX_TYPE),
    ],
)
def test_redteam_parser_bombs_are_rejected_without_growing_parent(path: Path, content_type: str):
    if not path.exists():
        pytest.skip("local red-team repro is not present")

    import resource
    from resume_parser import parse_resume_isolated

    rss_unit = 1 if sys.platform == "darwin" else 1024
    before_rss = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss * rss_unit
    started = time.monotonic()

    with pytest.raises(ValueError):
        parse_resume_isolated(path.name, content_type, path.read_bytes())

    elapsed = time.monotonic() - started
    after_rss = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss * rss_unit
    assert elapsed < 10
    assert after_rss - before_rss < 32 * 1024 * 1024
